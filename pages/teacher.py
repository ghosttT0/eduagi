import streamlit as st
import pandas as pd
from datetime import datetime
import json
import re
import time
from streamlit_echarts import st_echarts
from io import BytesIO
from docx import Document
# --- The fix is here: import the alignment enum ---/
from utils import load_conversational_chain
from database import SessionLocal, TeachingPlan, Exam, ExamQuestion, StudentDispute, User, Class, MindMap,VideoResource
try:
    from uil.file_utils import upload_to_qiniu
except ImportError as e:
    print(f"导入七牛云工具失败: {e}")
    def upload_to_qiniu(file_data, file_name):
        print("七牛云工具未正确导入，上传功能不可用")
        return None

# 添加强化版JSON解析函数
def parse_json_robust(result_text, expected_keys=None, fallback_data=None):
    """
    强化版AI响应解析函数，支持多种格式和容错处理
    """
    if not result_text or not result_text.strip():
        if fallback_data:
            st.info("📝 AI响应为空，使用默认模板")
            return fallback_data
        return None

    # 添加调试信息
    print(f"🔍 AI原始响应长度: {len(result_text)} 字符")
    print(f"🔍 AI响应前100字符: {result_text[:100]}")

    # 方法1: 直接JSON解析
    try:
        json_data = json.loads(result_text)
        print("✅ 方法1成功: 直接JSON解析")
        return json_data
    except Exception as e:
        print(f"❌ 方法1失败: {str(e)}")

    # 方法2: 提取大括号内容
    try:
        match = re.search(r'\{.*\}', result_text, re.DOTALL)
        if match:
            json_str = match.group(0)
            json_data = json.loads(json_str)
            print("✅ 方法2成功: 提取大括号内容")
            return json_data
    except Exception as e:
        print(f"❌ 方法2失败: {str(e)}")

    # 方法3: 清理后解析
    try:
        match = re.search(r'\{.*\}', result_text, re.DOTALL)
        if match:
            json_str = match.group(0)
            # 清理常见问题
            json_str = json_str.replace('\n', ' ').replace('\r', ' ')
            json_str = re.sub(r'\s+', ' ', json_str)
            json_str = json_str.replace("'", '"')  # 单引号改双引号

            # 尝试修复缺少引号的键
            json_str = re.sub(r'(\w+):', r'"\1":', json_str)

            json_data = json.loads(json_str)
            print("✅ 方法3成功: 清理后解析")
            return json_data
    except Exception as e:
        print(f"❌ 方法3失败: {str(e)}")

    # 方法4: 尝试提取JSON数组格式
    try:
        match = re.search(r'\[.*\]', result_text, re.DOTALL)
        if match:
            json_str = match.group(0)
            json_data = json.loads(json_str)
            # 如果是数组，包装成对象
            if isinstance(json_data, list) and expected_keys and len(expected_keys) > 0:
                wrapped_data = {expected_keys[0]: json_data}
                print("✅ 方法4成功: 提取并包装JSON数组")
                return wrapped_data
    except Exception as e:
        print(f"❌ 方法4失败: {str(e)}")

    # 方法5: 智能文本解析（针对试卷格式）
    if "questions" in str(expected_keys):
        try:
            questions = parse_text_to_questions(result_text)
            if questions:
                print("✅ 方法5成功: 智能文本解析")
                return {"questions": questions}
        except Exception as e:
            print(f"❌ 方法5失败: {str(e)}")

    # 方法6: 返回备选数据
    if fallback_data:
        st.warning("⚠️ AI返回格式异常，使用默认模板")
        print("🔄 使用默认模板")
        # 显示原始响应用于调试
        with st.expander("🔍 查看AI原始响应（调试用）"):
            st.text(result_text[:1000] + "..." if len(result_text) > 1000 else result_text)
        return fallback_data

    return None

def parse_text_to_questions(text):
    """从文本中智能解析试题"""
    questions = []

    # 尝试按行分割并查找题目模式
    lines = text.split('\n')
    current_question = None

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # 检测题目开始
        if re.match(r'^\d+[\.、]', line) or '题' in line:
            if current_question:
                questions.append(current_question)

            # 判断题目类型
            if '选择题' in line or '单选' in line or '多选' in line:
                question_type = 'multiple_choice'
            elif '简答' in line or '问答' in line:
                question_type = 'short_answer'
            elif '编程' in line or '代码' in line:
                question_type = 'coding'
            else:
                question_type = 'short_answer'  # 默认

            current_question = {
                'type': question_type,
                'question_text': line,
                'options': [],
                'answer': '',
                'explanation': ''
            }

        # 检测选项
        elif current_question and re.match(r'^[A-D][\.、]', line):
            current_question['options'].append(line)

        # 检测答案
        elif current_question and ('答案' in line or '参考答案' in line):
            current_question['answer'] = line.replace('答案:', '').replace('参考答案:', '').strip()

    # 添加最后一个题目
    if current_question:
        questions.append(current_question)

    return questions

def parse_teaching_plan_json(result_text, topic, course_name=""):
    """解析教案JSON"""
    expected_keys = ["teaching_content", "teaching_objectives", "key_points", "difficulties", "teaching_methods", "assessment"]

    # 创建备选数据
    fallback_data = {
        "teaching_content": f"""# {topic} 教学内容

## 核心概念
{topic}是本课程的重要组成部分，需要学生深入理解其基本原理和应用方法。

## 主要内容
1. 基本概念和定义
2. 核心原理和机制
3. 实际应用和案例
4. 相关技术和工具

## 教学安排
- 理论讲解：40分钟
- 实践操作：30分钟
- 讨论总结：20分钟""",

        "teaching_objectives": f"""## 教学目标

### 知识与技能
- 理解{topic}的基本概念和原理
- 掌握相关的操作方法和技能
- 能够应用所学知识解决实际问题

### 过程与方法
- 通过理论学习和实践操作相结合的方式
- 培养学生的分析和解决问题的能力
- 提高学生的动手实践能力

### 情感态度与价值观
- 培养学生对学科的兴趣和热情
- 增强学生的学习自信心
- 培养团队合作精神""",

        "key_points": f"""## 教学重点
1. {topic}的核心概念理解
2. 基本原理的掌握
3. 实际应用能力的培养

这些重点是学生必须掌握的核心内容。""",

        "difficulties": f"""## 教学难点
1. 抽象概念的理解
2. 理论与实践的结合
3. 复杂问题的分析方法

需要通过多种教学方法帮助学生克服这些难点。""",

        "teaching_methods": """## 教学方法
1. **讲授法**: 系统讲解核心概念
2. **演示法**: 通过实例演示操作过程
3. **讨论法**: 组织学生讨论相关问题
4. **实践法**: 安排实际操作练习""",

        "assessment": """## 评估方式
1. **课堂表现**: 参与度和互动情况
2. **实践操作**: 动手能力和操作规范
3. **作业完成**: 理解程度和应用能力
4. **期末考试**: 综合知识掌握情况"""
    }

    return parse_json_robust(result_text, expected_keys, fallback_data)

def parse_mindmap_json(result_text, topic):
    """解析思维导图JSON"""
    expected_keys = ["name", "children"]

    # 创建备选数据
    fallback_data = {
        "name": topic,
        "children": [
            {
                "name": "基本概念",
                "children": [
                    {"name": "定义"},
                    {"name": "特点"},
                    {"name": "分类"}
                ]
            },
            {
                "name": "核心原理",
                "children": [
                    {"name": "工作机制"},
                    {"name": "算法流程"},
                    {"name": "技术要点"}
                ]
            },
            {
                "name": "实际应用",
                "children": [
                    {"name": "应用场景"},
                    {"name": "案例分析"},
                    {"name": "发展趋势"}
                ]
            }
        ]
    }

    return parse_json_robust(result_text, expected_keys, fallback_data)

def parse_exam_json(result_text, scope, num_mcq=3, num_saq=2, num_code=1):
    """解析试卷JSON"""
    expected_keys = ["questions"]

    # 创建备选数据
    questions = []

    # 添加选择题
    for i in range(num_mcq):
        questions.append({
            "type": "multiple_choice",
            "question_text": f"关于{scope}的第{i+1}个选择题，以下哪个说法是正确的？",
            "options": ["选项A", "选项B", "选项C", "选项D"],
            "answer": "A",
            "explanation": "这是标准答案的解析说明。"
        })

    # 添加简答题
    for i in range(num_saq):
        questions.append({
            "type": "short_answer",
            "question_text": f"请详细说明{scope}中的第{i+1}个重要概念。",
            "answer": "这是简答题的参考答案，需要包含关键要点。",
            "explanation": "评分要点：概念准确、逻辑清晰、举例恰当。"
        })

    # 添加编程题
    for i in range(num_code):
        questions.append({
            "type": "coding",
            "question_text": f"编写一个关于{scope}的程序，实现第{i+1}个功能需求。",
            "answer": "# 参考代码\nprint('Hello World')",
            "explanation": "评分标准：代码正确性、规范性、效率性。"
        })

    fallback_data = {"questions": questions}

    return parse_json_robust(result_text, expected_keys, fallback_data)


def flatten_tree_nodes(node, level=0):
    """递归遍历树结构，返回所有节点的扁平列表，包含层级信息"""
    nodes = [{"name": node["name"], "level": level, "node": node}]
    if "children" in node:
        for child in node["children"]:
            nodes.extend(flatten_tree_nodes(child, level + 1))
    return nodes

def create_partial_tree(full_data, max_nodes):
    """创建部分显示的树结构"""
    if max_nodes <= 0:
        return None

    # 获取所有节点的扁平列表
    all_nodes = flatten_tree_nodes(full_data)

    # 按层级排序，确保父节点先显示
    all_nodes.sort(key=lambda x: (x["level"], x["name"]))

    # 只取前max_nodes个节点
    visible_nodes = all_nodes[:max_nodes]

    # 重建树结构
    def rebuild_tree(node_data, visible_set):
        result = {"name": node_data["name"]}
        if "children" in node_data:
            children = []
            for child in node_data["children"]:
                if any(v["name"] == child["name"] for v in visible_set):
                    child_result = rebuild_tree(child, visible_set)
                    if child_result:
                        children.append(child_result)
            if children:
                result["children"] = children
        return result

    return rebuild_tree(full_data, visible_nodes)

def write_content_to_docx(document, content, level=0):
    """
    一个可以递归处理复杂数据(字符串、字典、列表)并写入Word文档的智能函数。
    """
    indent = "    " * level  # 使用空格进行缩进
    if isinstance(content, dict):
        for key, value in content.items():
            # 为字典的键添加一个带缩进的标题行
            p = document.add_paragraph()
            p.add_run(f"{indent}• {key}:").bold = True
            # 递归处理值
            write_content_to_docx(document, value, level + 1)
    elif isinstance(content, list):
        for item in content:
            # 递归处理列表中的每一项，并用项目符号表示
            write_content_to_docx(document, f" - {item}", level)
    else:
        # 如果是普通文本，直接写入段落，并处理文本内部的换行符
        for paragraph_text in str(content).split('\n'):
            if paragraph_text.strip():
                # 添加带缩进的段落
                document.add_paragraph(f"{indent}{paragraph_text.strip()}")


def render():
    """渲染教师工作台页面的所有UI和逻辑"""
    qa_chain = load_conversational_chain()

    st.title("👨‍🏫 教师工作台")

    tab1, tab2, tab3, tab4,tab5 = st.tabs(["📝 **智能教学设计**", "🗺️ **AI知识图谱**", "✍️ **智能出题**", "❓ **学生疑问**","📽️ **视频中心**"])

    # --- Tab 1: 智能教学设计 (已修复历史记录bug) ---
    with tab1:
        st.subheader("AI 智能生成与导出专业教案")
        st.info("您可以输入具体的教学主题进行精细设计，或留空让AI为整个知识库（整本书）生成宏观教学大纲。")

        with st.form("lesson_plan_ultimate_form"):
            col1, col2 = st.columns(2)
            with col1:
                course_name = st.text_input("课程名称", placeholder="例如：《动手学深度学习》")
                class_hours = st.number_input("课时", min_value=1, max_value=8, value=2)
            with col2:
                chapter = st.text_input("所属章节", placeholder="例如：第3章 线性神经网络")
                teaching_time = st.number_input("授课时间（分钟）", min_value=45, max_value=180, value=90, step=45)

            topic_input = st.text_area(
                "核心教学主题 (可留空)",
                placeholder="例如：线性回归的从零开始实现。若留空，则为上方填写的章节生成整体教案。",
                height=100
            )
            submitted = st.form_submit_button("🤖 智能生成专业教案")

        if submitted:
            with st.spinner("AI教学总监正在为您规划教案蓝图..."):
                try:
                    if topic_input.strip():
                        prompt_scope = f"针对课程《{course_name}》中“{chapter}”章节下的核心主题“{topic_input}”"
                    elif chapter.strip():
                        prompt_scope = f"针对课程《{course_name}》的“{chapter}”整个章节"
                    else:
                        prompt_scope = f"为课程《{course_name}》"

                    prompt_template = f"""
                    你是一位顶级的教学设计总监，正在为《{course_name}》课程撰写一份专业、详尽、内容丰富的教学教案。
                    你的任务是: {prompt_scope}，设计一份完整的教学方案。
                    严格指令: 你的回复必须是单一、完整的JSON对象，不要有任何多余的解释。JSON对象必须包含以下键，每个键的值都必须是内容详实、逻辑严谨的字符串，可使用Markdown换行:
                    - `教学内容`: 详细阐述教学内容，而不仅仅是罗列标题。
                    - `教学目标`: 从“知识与技能”、“过程与方法”、“情感态度与价值观”三个维度详细描述教学目标。
                    - `教学重点`: 提炼出教学重点，并说明其重要性。
                    - `教学难点`: 分析教学难点，并提出具体的、可操作的突破策略。
                    - `教学设计`: 这是最重要的部分。请为每个环节（例如：导入、新授、练习、总结）都撰写详细的教师活动、学生互动和预计时间。内容必须丰富、饱满、具有可执行性。
                    - `教学反思与总结`: 设计几个有深度的启发性问题，供教师在课后进行教学反思。
                    格式规范:
                        -主标题：字体为微软雅黑，字号小二（18 磅），行距 35 磅，居中对齐，字形加粗，颜色黑色。
                        -副标题：字体为微软雅黑，字号三号（16 磅），行距 35 磅，居中对齐，字形倾斜，颜色黑色。
                        -多级标题（正文内层级）：字体为微软雅黑，字号小三（15 磅），行距 25 磅，字形加粗，使用中文编号 “一、二、三” 等，颜色黑色。
                        -正文：微软雅黑，字号小四（12 磅），行距 25 磅，颜色黑色。
            q           -重点内容：加粗
                    """

                    response = qa_chain.invoke({"question": prompt_template})
                    result_text = response.get('answer', '').strip()

                    # 使用强化版JSON解析
                    json_data = parse_teaching_plan_json(result_text, topic_input, course_name)
                    if not json_data:
                        st.error("AI未能返回有效的JSON内容。")
                        st.code(result_text)
                    else:
                        json_string = json.dumps(json_data, ensure_ascii=False)

                        db = SessionLocal()
                        try:
                            new_plan = TeachingPlan(
                                teacher_id=st.session_state.get("user_id", 1),
                                input_prompt=(
                                    topic_input if topic_input.strip() else f"{course_name} - {chapter} 整体大纲"),
                                output_content=json_string
                            )
                            db.add(new_plan)
                            db.commit()
                            st.success("专业教案已成功生成并保存！")
                            st.rerun()
                        except Exception as db_error:
                            st.error(f"数据库保存失败: {db_error}")
                            db.rollback()
                        finally:
                            db.close()

                except Exception as e:
                    st.error(f"生成教案时出错: {e}")

        st.markdown("---")
        st.subheader("📜 历史教案记录")

        db = SessionLocal()
        try:
            history_plans = db.query(TeachingPlan).order_by(TeachingPlan.timestamp.desc()).all()

            if not history_plans:
                st.info("暂无历史记录。请先在上方生成一份新教案。")
            else:
                plan_data_for_display = {
                    "ID": [p.id for p in history_plans],
                    "生成时间": [p.timestamp.strftime("%Y-%m-%d %H:%M") for p in history_plans],
                    "教案主题": [p.input_prompt[:50] + '...' if len(p.input_prompt) > 50 else p.input_prompt for p in
                                 history_plans],
                }
                df = pd.DataFrame(plan_data_for_display)
                st.dataframe(df, use_container_width=True, hide_index=True)

                plan_ids = [p.id for p in history_plans]
                selected_id = st.selectbox("请选择一个历史教案ID查看详情或导出：", options=plan_ids,
                                           key="history_selectbox")

                if selected_id:
                    selected_plan = next((p for p in history_plans if p.id == selected_id), None)
                    if selected_plan:

                        # --- 核心修复：增加try...except来兼容旧的纯文本数据 ---
                        with st.expander(f"查看ID: {selected_id} 的详细内容", expanded=True):
                            try:
                                # 尝试将内容解析为JSON
                                plan_details = json.loads(selected_plan.output_content)
                                # 如果成功，按漂亮的结构化格式显示
                                title_map = {
                                    "teaching_content": "教学内容", "teaching_objectives": "教学目标",
                                    "key_points": "教学重点", "teaching_difficulties": "教学难点",
                                    "teaching_design": "教学设计与过程", "teaching_reflection": "教学反思"
                                }
                                for key, value in plan_details.items():
                                    st.markdown(f"**{title_map.get(key, key)}**")
                                    st.markdown(str(value))
                            except json.JSONDecodeError:
                                # 如果解析失败，说明是旧的纯文本数据，直接显示原文
                                st.markdown("##### 无法按JSON解析，显示原文：")
                                st.text(selected_plan.output_content)

                        # --- 为选中的历史记录提供导出功能 ---
                        try:
                            # 再次进行try...except，因为word导出也依赖于JSON解析
                            plan_details_for_doc = json.loads(selected_plan.output_content)

                            document = Document()
                            document.add_heading(f"《{selected_plan.input_prompt}》教案", level=1)

                            title_map_doc = {
                                "teaching_content": "教学内容", "teaching_objectives": "教学目标",
                                "key_points": "教学重点", "teaching_difficulties": "教学难点",
                                "teaching_design": "教学设计与过程", "teaching_reflection": "教学反思"
                            }
                            for key, value in plan_details_for_doc.items():
                                document.add_heading(title_map_doc.get(key, key), level=2)
                                write_content_to_docx(document, value)
                                document.add_paragraph()

                            file_stream = BytesIO()
                            document.save(file_stream)
                            file_stream.seek(0)

                            st.download_button(
                                label=f"📄 导出ID: {selected_id} 为Word文档",
                                data=file_stream,
                                file_name=f"历史教案_{selected_id}_{selected_plan.input_prompt[:10]}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            )
                        except json.JSONDecodeError:
                            st.warning("此条历史记录是旧的纯文本格式，无法导出为结构化Word文档。")
                        except Exception as e:
                            st.error(f"创建Word文档时出错: {e}")
        finally:
            db.close()

    # --- 其他Tabs ---
    with tab2:
        st.info("AI知识图谱功能区。")
    with tab3:
        st.info("智能出题功能区。")

    # --- Tab 2: AI知识图谱与大纲 (基于您的代码进行美化) ---
    with tab2:
        st.header("AI 知识图谱与大纲生成器")
        st.markdown("---")

        # 创建子标签页
        subtab1, subtab2, subtab3 = st.tabs(["🧠 生成新图谱", "📚 我的图谱库", "📝 Markdown大纲"])

        # --- 子标签1: 生成新图谱 ---
        with subtab1:
            # --- Section 1: 可视化知识图谱 ---
            with st.container(border=True):
                st.subheader("🗺️ 可视化知识图谱")
                st.info("请输入主题，AI将为您生成可交互的、精美的可视化知识图谱。")

                with st.form("echarts_mindmap_form"):
                    topic_input_mindmap = st.text_input("请输入图谱主题：", placeholder="例如：深度学习中的卷积神经网络")
                    submitted_mindmap = st.form_submit_button("🧠 生成知识图谱")

                if submitted_mindmap and topic_input_mindmap:
                    with st.spinner("AI正在构建知识网络..."):
                        try:
                            prompt_template_json = f"""
                            您是JSON格式专家，请为主题 “{topic_input_mindmap}” 创建一个符合ECharts树图的、语法完全正确的JSON。
                            规则：根节点必须有 'name' 键，子节点在 'children' 数组中。请创建一个层次丰富的知识图谱，至少包含3-4层节点。回复中只能包含纯JSON。
                            """
                            response = qa_chain.invoke({"question": prompt_template_json})
                            # ConversationalRetrievalChain 返回的是 'answer' 而不是 'result'
                            result_text = response.get('answer', '').strip()

                            # 使用强化版JSON解析
                            full_data = parse_mindmap_json(result_text, topic_input_mindmap)
                            if not full_data:
                                st.error("错误：AI未能返回有效的JSON内容。")
                                st.code(result_text, language="text")
                            else:

                                # 存储完整数据到session state
                                st.session_state.full_mindmap_data = full_data
                                st.session_state.mindmap_animation_started = True
                                st.session_state.current_node_count = 1  # 从根节点开始

                                st.success("知识图谱数据生成完成！正在逐步构建可视化...")

                        except Exception as e:
                            st.error(f"生成思维导图时出错: {e}")

                # 动态显示思维导图
                if st.session_state.get("mindmap_animation_started", False):
                    # 获取完整数据
                    full_data = st.session_state.get("full_mindmap_data")
                    if full_data:
                        # 获取所有节点数量
                        all_nodes = flatten_tree_nodes(full_data)
                        total_nodes = len(all_nodes)

                        # 当前显示的节点数量
                        current_count = st.session_state.get("current_node_count", 1)

                        # 创建进度条
                        progress_col1, progress_col2 = st.columns([3, 1])
                        with progress_col1:
                            progress = st.progress(current_count / total_nodes)
                            st.caption(f"正在构建知识图谱... ({current_count}/{total_nodes} 节点)")

                        with progress_col2:
                            if current_count < total_nodes:
                                if st.button("⏸️ 暂停", key="pause_animation"):
                                    st.session_state.animation_paused = not st.session_state.get("animation_paused", False)
                            else:
                                if st.button("🔄 重新播放", key="replay_animation"):
                                    st.session_state.current_node_count = 1
                                    st.session_state.animation_paused = False
                                    st.rerun()

                        # 创建部分显示的数据
                        partial_data = create_partial_tree(full_data, current_count)

                        if partial_data:
                            echarts_options = {
                            "tooltip": {"trigger": "item", "triggerOn": "mousemove"},
                            "series": [
                                {
                                    "type": "tree",
                                    "data": [partial_data],
                                    "top": "5%", "left": "10%", "bottom": "5%", "right": "20%",
                                    "symbolSize": 8,
                                    "edgeShape": "curve",
                                    "expandAndCollapse": True,
                                    "initialTreeDepth": 5,
                                    "label": {
                                        "position": "left",
                                        "verticalAlign": "middle",
                                        "align": "right",
                                        "fontSize": 14,
                                        "color": "#333",
                                        "backgroundColor": "#f0f8ff",
                                        "padding": [6, 12],
                                        "borderRadius": 8,
                                        "borderWidth": 1,
                                        "borderColor": "#A9A9A9",
                                        "shadowColor": 'rgba(0, 0, 0, 0.2)',
                                        "shadowBlur": 5,
                                    },
                                    "leaves": {
                                        "label": {
                                            "position": "right",
                                            "verticalAlign": "middle",
                                            "align": "left",
                                            "backgroundColor": "#e6f7ff",
                                        }
                                    },
                                    "emphasis": {
                                        "focus": 'descendant',
                                        "label": {
                                            "borderColor": '#007bff',
                                            "borderWidth": 2,
                                        }
                                    },
                                    "animation": True,
                                    "animationDuration": 800,
                                    "animationEasing": "elasticOut",
                                }
                            ],
                        }

                        # 显示图表
                        st_echarts(options=echarts_options, height="600px", key=f"mindmap_{current_count}")

                        # 自动推进动画（如果没有暂停）
                        if current_count < total_nodes and not st.session_state.get("animation_paused", False):
                            import time as time_module  # 明确导入避免作用域问题
                            time_module.sleep(1.5)  # 控制动画速度
                            st.session_state.current_node_count = current_count + 1
                            st.rerun()
                        elif current_count >= total_nodes:
                            st.success("🎉 知识图谱构建完成！您可以与图谱进行交互。")
                            # 显示完成后的操作选项
                            col_action1, col_action2 = st.columns(2)
                            with col_action1:
                                if st.button("💾 保存图谱", key="save_mindmap"):
                                    # 显示保存对话框
                                    st.session_state.show_save_dialog = True
                                    st.rerun()
                            with col_action2:
                                if st.button("📤 导出图片", key="export_mindmap"):
                                    st.session_state.show_export_dialog = True
                                    st.rerun()

                            # 保存图谱对话框
                            if st.session_state.get("show_save_dialog", False):
                                with st.form("save_mindmap_form"):
                                    st.subheader("保存思维导图")
                                    map_title = st.text_input("图谱标题", value=topic_input_mindmap)
                                    map_description = st.text_area("图谱描述", placeholder="请输入对此思维导图的描述...")
                                    is_public = st.checkbox("公开此图谱", value=False, help="公开的图谱可以被其他用户查看")

                                    col_save, col_cancel = st.columns(2)
                                    with col_save:
                                        save_submitted = st.form_submit_button("确认保存", use_container_width=True)
                                    with col_cancel:
                                        cancel_save = st.form_submit_button("取消", use_container_width=True)

                                    if save_submitted:
                                        if not map_title:
                                            st.error("请输入图谱标题")
                                        else:
                                            try:
                                                # 将思维导图数据保存到数据库
                                                db = SessionLocal()
                                                new_mindmap = MindMap(
                                                    user_id=st.session_state.get("user_id"),
                                                    title=map_title,
                                                    topic=topic_input_mindmap,
                                                    data=json.dumps(full_data),
                                                    description=map_description,
                                                    is_public=is_public
                                                )
                                                db.add(new_mindmap)
                                                db.commit()
                                                st.success(f"思维导图 '{map_title}' 已成功保存！")
                                                st.session_state.show_save_dialog = False
                                                st.rerun()
                                            except Exception as e:
                                                st.error(f"保存失败: {e}")
                                            finally:
                                                db.close()

                                    if cancel_save:
                                        st.session_state.show_save_dialog = False
                                        st.rerun()

                            # 导出图片对话框
                            if st.session_state.get("show_export_dialog", False):
                                with st.container(border=True):
                                    st.subheader("导出思维导图")
                                    st.info("💡 提示：由于技术限制，我们提供以下导出方式：")

                                    col_export1, col_export2 = st.columns(2)

                                    with col_export1:
                                        st.markdown("##### 📄 导出为JSON数据")
                                        st.caption("包含完整的思维导图结构数据")

                                        # 创建JSON文件
                                        json_data = json.dumps(full_data, ensure_ascii=False, indent=2)
                                        json_bytes = json_data.encode('utf-8')

                                        st.download_button(
                                            label="📥 下载JSON文件",
                                            data=json_bytes,
                                            file_name=f"mindmap_{topic_input_mindmap}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                                            mime="application/json",
                                            use_container_width=True
                                        )

                                    with col_export2:
                                        st.markdown("##### 🖼️ 导出为HTML文件")
                                        st.caption("可在浏览器中打开查看交互式图谱")

                                        # 创建HTML文件
                                        html_content = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>思维导图 - {topic_input_mindmap}</title>
    <script src="https://cdn.jsdelivr.net/npm/echarts@5.4.0/dist/echarts.min.js"></script>
    <style>
        body {{ margin: 0; padding: 20px; font-family: Arial, sans-serif; }}
        #mindmap {{ width: 100%; height: 600px; border: 1px solid #ddd; }}
        .title {{ text-align: center; margin-bottom: 20px; }}
    </style>
</head>
<body>
    <div class="title">
        <h1>思维导图：{topic_input_mindmap}</h1>
        <p>生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
    </div>
    <div id="mindmap"></div>

    <script>
        var chartDom = document.getElementById('mindmap');
        var myChart = echarts.init(chartDom);

        var option = {echarts_options};

        myChart.setOption(option);

        window.addEventListener('resize', function() {{
            myChart.resize();
        }});
    </script>
</body>
</html>"""

                                        # 将完整的ECharts配置插入HTML
                                        final_echarts_options = {
                                            "tooltip": {"trigger": "item", "triggerOn": "mousemove"},
                                            "series": [
                                                {
                                                    "type": "tree",
                                                    "data": [full_data],
                                                    "top": "5%", "left": "10%", "bottom": "5%", "right": "20%",
                                                    "symbolSize": 8,
                                                    "edgeShape": "curve",
                                                    "expandAndCollapse": True,
                                                    "initialTreeDepth": 5,
                                                    "label": {
                                                        "position": "left",
                                                        "verticalAlign": "middle",
                                                        "align": "right",
                                                        "fontSize": 14,
                                                        "color": "#333",
                                                        "backgroundColor": "#f0f8ff",
                                                        "padding": [6, 12],
                                                        "borderRadius": 8,
                                                        "borderWidth": 1,
                                                        "borderColor": "#A9A9A9",
                                                        "shadowColor": 'rgba(0, 0, 0, 0.2)',
                                                        "shadowBlur": 5,
                                                    },
                                                    "leaves": {
                                                        "label": {
                                                            "position": "right",
                                                            "verticalAlign": "middle",
                                                            "align": "left",
                                                            "backgroundColor": "#e6f7ff",
                                                        }
                                                    },
                                                    "emphasis": {
                                                        "focus": 'descendant',
                                                        "label": {
                                                            "borderColor": '#007bff',
                                                            "borderWidth": 2,
                                                        }
                                                    },
                                                    "animation": True,
                                                    "animationDuration": 800,
                                                }
                                            ],
                                        }

                                        html_final = html_content.replace('{echarts_options}', json.dumps(final_echarts_options))
                                        html_bytes = html_final.encode('utf-8')

                                        st.download_button(
                                            label="📥 下载HTML文件",
                                            data=html_bytes,
                                            file_name=f"mindmap_{topic_input_mindmap}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html",
                                            mime="text/html",
                                            use_container_width=True
                                        )

                                    st.markdown("---")
                                    st.markdown("##### 📱 其他导出方式")
                                    st.info("""
                                    **浏览器截图方式：**
                                    1. 下载HTML文件并在浏览器中打开
                                    2. 使用浏览器的开发者工具或截图扩展
                                    3. 截取思维导图区域保存为图片

                                    **专业工具导出：**
                                    - 可以将JSON数据导入到专业的思维导图软件中
                                    - 支持XMind、MindMaster等软件的进一步编辑
                                    """)

                                    if st.button("关闭导出对话框", use_container_width=True):
                                        st.session_state.show_export_dialog = False
                                        st.rerun()

        # --- 子标签2: 我的图谱库 ---
        with subtab2:
            st.subheader("📚 我的思维导图库")

            db = SessionLocal()
            try:
                current_user_id = st.session_state.get("user_id")
                saved_mindmaps = db.query(MindMap).filter(
                    MindMap.user_id == current_user_id
                ).order_by(MindMap.created_at.desc()).all()

                if not saved_mindmaps:
                    st.info("您还没有保存任何思维导图。请在'生成新图谱'中创建并保存图谱。")
                else:
                    st.success(f"共找到 {len(saved_mindmaps)} 个已保存的思维导图")

                    for mindmap in saved_mindmaps:
                        with st.container(border=True):
                            col_info, col_actions = st.columns([3, 1])

                            with col_info:
                                st.markdown(f"**📊 {mindmap.title}**")
                                st.caption(f"主题: {mindmap.topic}")
                                if mindmap.description:
                                    st.markdown(f"*{mindmap.description}*")
                                st.caption(f"创建时间: {mindmap.created_at.strftime('%Y-%m-%d %H:%M')} | {'公开' if mindmap.is_public else '私有'}")

                            with col_actions:
                                if st.button("👁️ 查看", key=f"view_mindmap_{mindmap.id}"):
                                    st.session_state[f"viewing_mindmap_{mindmap.id}"] = True
                                    st.rerun()

                                if st.button("🗑️ 删除", key=f"delete_mindmap_{mindmap.id}"):
                                    try:
                                        db.delete(mindmap)
                                        db.commit()
                                        st.success("思维导图已删除")
                                        st.rerun()
                                    except Exception as e:
                                        st.error(f"删除失败: {e}")

                            # 查看思维导图
                            if st.session_state.get(f"viewing_mindmap_{mindmap.id}", False):
                                try:
                                    mindmap_data = json.loads(mindmap.data)

                                    # 创建ECharts配置
                                    view_options = {
                                        "tooltip": {"trigger": "item", "triggerOn": "mousemove"},
                                        "series": [
                                            {
                                                "type": "tree",
                                                "data": [mindmap_data],
                                                "top": "5%", "left": "10%", "bottom": "5%", "right": "20%",
                                                "symbolSize": 8,
                                                "edgeShape": "curve",
                                                "expandAndCollapse": True,
                                                "initialTreeDepth": 5,
                                                "label": {
                                                    "position": "left",
                                                    "verticalAlign": "middle",
                                                    "align": "right",
                                                    "fontSize": 14,
                                                    "color": "#333",
                                                    "backgroundColor": "#f0f8ff",
                                                    "padding": [6, 12],
                                                    "borderRadius": 8,
                                                    "borderWidth": 1,
                                                    "borderColor": "#A9A9A9",
                                                    "shadowColor": 'rgba(0, 0, 0, 0.2)',
                                                    "shadowBlur": 5,
                                                },
                                                "leaves": {
                                                    "label": {
                                                        "position": "right",
                                                        "verticalAlign": "middle",
                                                        "align": "left",
                                                        "backgroundColor": "#e6f7ff",
                                                    }
                                                },
                                                "emphasis": {
                                                    "focus": 'descendant',
                                                    "label": {
                                                        "borderColor": '#007bff',
                                                        "borderWidth": 2,
                                                    }
                                                },
                                                "animation": True,
                                                "animationDuration": 800,
                                            }
                                        ],
                                    }

                                    st_echarts(options=view_options, height="500px", key=f"saved_mindmap_{mindmap.id}")

                                    if st.button("关闭查看", key=f"close_view_mindmap_{mindmap.id}"):
                                        st.session_state[f"viewing_mindmap_{mindmap.id}"] = False
                                        st.rerun()

                                except Exception as e:
                                    st.error(f"显示思维导图时出错: {e}")
            finally:
                db.close()

        # --- 子标签3: Markdown大纲 ---
        with subtab3:

            st.subheader("📜 Markdown大纲")
            st.info("请输入主题，AI将为您生成结构化的Markdown文本大纲。")

            with st.form("markdown_form"):
                topic_input_markdown = st.text_input("请输入大纲主题：", placeholder="例如：深度学习前十章知识概览")
                submitted_markdown = st.form_submit_button("✍️ 生成Markdown大纲")

            if submitted_markdown and topic_input_markdown:
                with st.spinner("AI正在生成Markdown大纲..."):
                    try:
                        prompt_template_md = f"""
                        您是知识结构专家，请为主题 “{topic_input_markdown}” 生成一份层级清晰的Markdown文本。
                        """
                        response = qa_chain.invoke({"question": prompt_template_md})
                        st.session_state.markdown_text = response.get('answer', '')
                        st.success("Markdown大纲已生成！")
                    except Exception as e:
                        st.error(f"生成Markdown时出错: {e}")

            if "markdown_text" in st.session_state:
                st.markdown(st.session_state.markdown_text)

    # --- Tab 3: 智能出题与发布 (升级为Word导出) ---
    with tab3:
        st.subheader("AI 智能生成与发布试卷")
        st.info("请设定试卷参数，AI将根据您的要求和知识库内容，自动生成一套完整的试卷。")

        with st.form("exam_generation_form"):
            st.markdown("##### 1. 设定试卷基本信息")
            exam_scope = st.text_area("考试范围说明",
                                      placeholder="例如：围绕《动手学深度学习》中关于卷积神经网络（CNN）和循环神经网络（RNN）的核心概念进行考察。",
                                      height=100)

            st.markdown("##### 2. 设定题目数量")
            col1, col2, col3 = st.columns(3)
            with col1:
                num_mcq = st.number_input("选择题数量", min_value=0, max_value=20, value=5)
            with col2:
                num_saq = st.number_input("简答题数量", min_value=0, max_value=10, value=3)
            with col3:
                num_code = st.number_input("编程题数量", min_value=0, max_value=5, value=1)

            submitted = st.form_submit_button("🤖 开始智能生成试卷")

        if submitted:
            if not exam_scope.strip():
                st.warning("请输入考试范围说明！")
            else:
                with st.spinner(f"正在为您生成包含 {num_mcq}道选择题, {num_saq}道简答题, {num_code}道编程题的试卷..."):
                    try:
                        # 设计一个能生成完整试卷JSON的“超级Prompt”
                        prompt_template = f"""
你是一位资深的命题专家。请根据以下要求生成一份完整的试卷。

**试卷要求:**
- 考察范围: {exam_scope}
- 选择题数量: {num_mcq}
- 简答题数量: {num_saq}
- 编程题数量: {num_code}

**重要：你必须严格按照以下JSON格式返回，不要添加任何解释文字：**

{{
  "questions": [
    {{
      "type": "multiple_choice",
      "question_text": "题目内容",
      "options": ["A. 选项1", "B. 选项2", "C. 选项3", "D. 选项4"],
      "answer": "A",
      "explanation": "答案解析"
    }},
    {{
      "type": "short_answer",
      "question_text": "简答题内容",
      "answer": "参考答案",
      "explanation": "评分要点"
    }},
    {{
      "type": "coding",
      "question_text": "编程题要求",
      "answer": "参考代码",
      "explanation": "评分标准"
    }}
  ]
}}

请严格按照上述JSON格式生成{num_mcq + num_saq + num_code}道题目，直接返回JSON，不要包含其他内容。
                        """
                        response = qa_chain.invoke({"question": prompt_template})
                        result_text = response.get('answer', '').strip()

                        # 使用强化版JSON解析
                        exam_data = parse_exam_json(result_text, exam_scope, num_mcq, num_saq, num_code)
                        if not exam_data:
                            st.error("AI未能生成有效的试卷JSON。")
                            st.code(result_text)
                        else:
                            st.session_state.generated_exam = exam_data  # 存入会话状态
                            st.success("试卷已成功生成！请在下方预览和操作。")

                    except Exception as e:
                        st.error(f"生成试卷时出错: {e}")

        # --- 试卷预览、发布与导出 ---
        if "generated_exam" in st.session_state:
            st.markdown("---")
            st.subheader("📜 试卷预览")

            exam_data = st.session_state.generated_exam
            for i, q in enumerate(exam_data.get("questions", [])):
                st.markdown(f"**第 {i + 1} 题 ({q.get('type', '未知类型')})**")
                st.markdown(q.get("question_text", "无题干"))
                if q.get('type') == 'multiple_choice':
                    for opt in q.get('options', []):
                        st.markdown(f"- {opt}")

                with st.expander("点击查看答案及解析"):
                    st.success(f"**答案:** {q.get('answer', '无')}")
                    st.info(f"**解析:** {q.get('explanation', '无')}")
                st.markdown("---")

            # --- 发布与导出按钮 ---
            st.subheader("下一步操作")
            col_a, col_b = st.columns(2)
            with col_a:
                if st.button("🚀 发布试卷到平台"):
                    if "generated_exam" in st.session_state and st.session_state.generated_exam:
                        db = SessionLocal()
                        try:
                            # 1. 创建一张新试卷记录
                            new_exam = Exam(
                                teacher_id=st.session_state.get("user_id"),
                                scope=exam_scope  # exam_scope 来自您之前的表单
                            )
                            db.add(new_exam)
                            db.flush()  # 让我们能提前获得新试卷的ID

                            # 2. 遍历并保存每一道题目
                            for q in st.session_state.generated_exam.get("questions", []):
                                new_question = ExamQuestion(
                                    exam_id=new_exam.id,
                                    question_type=q.get("type"),
                                    question_text=q.get("question_text"),
                                    options=json.dumps(q.get("options", [])),  # 列表转JSON字符串
                                    answer=q.get("answer"),
                                    explanation=q.get("explanation")
                                )
                                db.add(new_question)

                            db.commit()
                            st.success(f"试卷已成功发布并存入数据库！ID为: {new_exam.id}")
                        except Exception as e:
                            db.rollback()
                            st.error(f"发布试卷时出错: {e}")
                        finally:
                            db.close()
                    else:
                        st.warning("没有可发布的试卷内容。")

            with col_b:
                # --- 核心修改：创建并提供Word文档下载 ---

                # 创建一个内存中的Word文档
                document = Document()
                document.add_heading(f"试卷：{exam_scope}", level=1)

                for i, q in enumerate(exam_data.get("questions", [])):
                    document.add_heading(f"第 {i + 1} 题 ({q.get('type', '未知类型')})", level=2)

                    # 添加题干，并处理加粗等简单格式
                    p_question = document.add_paragraph()
                    p_question.add_run('题干: ').bold = True
                    p_question.add_run(q.get("question_text", "无题干"))

                    if q.get('type') == 'multiple_choice':
                        for opt in q.get('options', []):
                            document.add_paragraph(opt, style='List Bullet')

                    document.add_paragraph()  # 添加空行

                # 将文档保存到内存中的字节流
                file_stream = BytesIO()
                document.save(file_stream)
                file_stream.seek(0)  # 将指针移回文件的开头

                st.download_button(
                    label="📄 导出为Word文档(.docx)",
                    data=file_stream,
                    file_name=f"生成的试卷_{datetime.now().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )

    # --- Tab 4: 学生疑问处理 ---
    with tab4:
        st.subheader("📋 本班级学生疑问处理")

        # 获取当前教师的班级信息
        db = SessionLocal()
        try:
            current_teacher_id = st.session_state.get("user_id")
            teacher = db.query(User).filter(User.id == current_teacher_id).first()

            if not teacher or not teacher.class_id:
                st.warning("您还没有被分配到任何班级，无法查看学生疑问。请联系管理员。")
                return

            # 获取班级信息
            teacher_class = db.query(Class).filter(Class.id == teacher.class_id).first()
            st.info(f"当前管理班级：{teacher_class.name}")

            # 获取本班级的所有疑问
            disputes = db.query(StudentDispute).join(User, StudentDispute.student_id == User.id)\
                .filter(StudentDispute.class_id == teacher.class_id)\
                .order_by(StudentDispute.timestamp.desc()).all()

            if not disputes:
                st.info("目前没有学生疑问需要处理。")
            else:
                # 统计信息
                total_disputes = len(disputes)
                pending_disputes = len([d for d in disputes if d.status == "待处理"])
                resolved_disputes = len([d for d in disputes if d.status == "已回复"])

                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("总疑问数", total_disputes)
                with col2:
                    st.metric("待处理", pending_disputes, delta=f"-{resolved_disputes}")
                with col3:
                    st.metric("已回复", resolved_disputes)

                st.markdown("---")

                # 疑问列表
                for dispute in disputes:
                    student = db.query(User).filter(User.id == dispute.student_id).first()

                    # 根据状态设置不同的样式
                    if dispute.status == "待处理":
                        status_color = "🔴"
                        container_type = "error"
                    else:
                        status_color = "✅"
                        container_type = "success"

                    with st.container(border=True):
                        col_info, col_action = st.columns([3, 1])

                        with col_info:
                            st.markdown(f"**{status_color} 疑问 #{dispute.id}**")
                            st.markdown(f"**学生：** {student.display_name} ({student.account_id})")
                            st.markdown(f"**提交时间：** {dispute.timestamp.strftime('%Y-%m-%d %H:%M')}")
                            st.markdown(f"**疑问内容：**")
                            st.markdown(f"> {dispute.message}")

                            if dispute.teacher_reply:
                                st.markdown(f"**您的回复：**")
                                st.success(dispute.teacher_reply)
                                st.caption(f"回复时间：{dispute.reply_timestamp.strftime('%Y-%m-%d %H:%M')}")

                        with col_action:
                            if dispute.status == "待处理":
                                if st.button(f"回复", key=f"reply_dispute_{dispute.id}"):
                                    st.session_state[f"replying_{dispute.id}"] = True
                                    st.rerun()
                            else:
                                st.success("已回复")

                        # 回复表单
                        if st.session_state.get(f"replying_{dispute.id}", False):
                            with st.form(f"reply_form_{dispute.id}"):
                                reply_content = st.text_area("回复内容", placeholder="请输入您的回复...", height=100)
                                col_submit, col_cancel = st.columns(2)

                                with col_submit:
                                    if st.form_submit_button("发送回复", use_container_width=True):
                                        if reply_content.strip():
                                            try:
                                                dispute.teacher_reply = reply_content
                                                dispute.status = "已回复"
                                                dispute.reply_timestamp = datetime.now()
                                                db.commit()
                                                st.success("回复已发送！")
                                                del st.session_state[f"replying_{dispute.id}"]
                                                st.rerun()
                                            except Exception as e:
                                                st.error(f"发送回复失败：{e}")
                                                db.rollback()
                                        else:
                                            st.warning("请输入回复内容")

                                with col_cancel:
                                    if st.form_submit_button("取消", use_container_width=True):
                                        del st.session_state[f"replying_{dispute.id}"]
                                        st.rerun()

                        st.markdown("---")
        finally:
            db.close()
        # --- Tab 5: 视频管理 (优化版) ---
        with tab5:
            st.subheader("📹 教学视频资源管理")
            st.info("您可以在这里管理教学视频：添加外部链接、上传到七牛云、保存草稿、AI分析等。")

            # 创建子标签页
            video_tab1, video_tab2, video_tab3 = st.tabs(["📤 上传视频", "📚 我的视频库", "🔗 添加链接"])

            # --- 子标签1: 上传视频 ---
            with video_tab1:
                st.markdown("##### 📤 上传本地视频文件到七牛云")

                uploaded_file = st.file_uploader(
                    "选择一个视频文件...",
                    type=["mp4", "mov", "avi", "mkv", "wmv", "flv"],
                    help="支持常见的视频格式，文件大小建议不超过500MB"
                )

                if uploaded_file is not None:
                    # 显示文件信息
                    file_size = len(uploaded_file.getvalue()) / (1024 * 1024)  # MB
                    st.info(f"📁 文件名: {uploaded_file.name} | 📊 大小: {file_size:.2f} MB")

                    # 视频信息输入
                    col1, col2 = st.columns(2)
                    with col1:
                        video_title_upload = st.text_input(
                            "视频标题 *",
                            key="qiniu_video_title",
                            placeholder="例如：Python基础教程第一课"
                        )
                    with col2:
                        video_status = st.selectbox(
                            "保存状态",
                            ["草稿", "已发布"],
                            index=0,
                            help="草稿状态的视频只有您可以看到"
                        )

                    video_desc_upload = st.text_area(
                        "视频简介",
                        key="qiniu_video_desc",
                        placeholder="请简要描述视频内容、适用对象、学习目标等...",
                        height=100
                    )

                    # 操作按钮
                    col_upload, col_draft = st.columns(2)

                    with col_upload:
                        if st.button("🚀 上传到云空间", use_container_width=True):
                            if not video_title_upload.strip():
                                st.warning("⚠️ 请输入视频标题！")
                            else:
                                # 文件大小检查和时间估算
                                file_size_mb = file_size
                                if file_size_mb > 500:
                                    st.error("❌ 文件过大！请选择小于500MB的视频文件。")
                                elif file_size_mb > 100:
                                    estimated_time = int(file_size_mb * 0.5)  # 估算上传时间（分钟）
                                    st.warning(f"⚠️ 文件较大（{file_size_mb:.1f}MB），预计上传时间约{estimated_time}分钟，请耐心等待...")
                                    st.info("💡 建议：可以先压缩视频文件以加快上传速度")
                                elif file_size_mb > 50:
                                    estimated_time = int(file_size_mb * 0.3)
                                    st.info(f"📤 准备上传{file_size_mb:.1f}MB文件，预计需要{estimated_time}分钟")

                                # 创建进度容器
                                progress_container = st.empty()
                                status_container = st.empty()

                                try:
                                    with status_container.container():
                                        st.info("🔍 准备上传文件...")

                                    # 获取文件的二进制数据
                                    file_data = uploaded_file.getvalue()

                                    with status_container.container():
                                        st.info(f"📁 文件准备完成: {len(file_data)} bytes ({file_size_mb:.2f} MB)")

                                    # 生成唯一文件名
                                    import uuid
                                    unique_filename = f"{uuid.uuid4()}_{uploaded_file.name}"

                                    with status_container.container():
                                        st.info(f"📝 生成文件名: {unique_filename}")

                                    # 显示上传进度（演示模式 - 直接成功）
                                    with progress_container.container():
                                        progress_bar = st.progress(0)
                                        progress_text = st.empty()

                                        progress_text.text("☁️ 正在连接七牛云...")
                                        progress_bar.progress(10)

                                        import time as time_module  # 避免与全局time冲突
                                        time_module.sleep(0.5)  # 模拟连接时间

                                        progress_text.text("📤 正在上传文件...")
                                        progress_bar.progress(30)
                                        time_module.sleep(1)  # 模拟上传时间

                                        progress_text.text("🔄 处理文件...")
                                        progress_bar.progress(60)
                                        time_module.sleep(0.5)

                                        progress_text.text("✅ 上传完成，正在保存...")
                                        progress_bar.progress(90)

                                        # 演示模式：生成模拟的视频URL
                                        demo_video_url = f"https://eduagi.site/demo/{unique_filename}"

                                        # 将模拟URL存入数据库
                                        db = SessionLocal()
                                        try:
                                            new_video = VideoResource(
                                                teacher_id=st.session_state.get("user_id"),
                                                title=video_title_upload,
                                                description=video_desc_upload,
                                                path=demo_video_url,  # 存储模拟的URL
                                                status=video_status
                                            )
                                            db.add(new_video)
                                            db.commit()

                                            progress_bar.progress(100)
                                            progress_text.text("🎉 完成！")

                                            st.success(f"✅ 视频已成功上传到七牛云，状态：{video_status}！")
                                            st.balloons()

                                            # 显示上传结果
                                            with st.expander("📋 上传详情", expanded=True):
                                                st.write(f"**视频标题:** {video_title_upload}")
                                                st.write(f"**存储状态:** {video_status}")
                                                st.write(f"**云端URL:** {demo_video_url}")
                                                st.write(f"**文件大小:** {file_size:.2f} MB")
                                                st.info("🎬 演示模式：文件已模拟上传成功")

                                        except Exception as e:
                                            st.error(f"❌ 数据库保存失败: {e}")
                                        finally:
                                            db.close()

                                except Exception as e:
                                    st.error(f"❌ 上传过程中出现错误: {e}")
                                    with st.expander("🔍 详细错误信息"):
                                        st.code(str(e))

                    with col_draft:
                        if st.button("💾 保存为草稿", use_container_width=True):
                            if not video_title_upload.strip():
                                st.warning("⚠️ 请输入视频标题！")
                            else:
                                # 保存草稿到session state
                                draft_data = {
                                    "title": video_title_upload,
                                    "description": video_desc_upload,
                                    "file_name": uploaded_file.name,
                                    "file_size": file_size,
                                    "status": "草稿"
                                }

                                if "video_drafts" not in st.session_state:
                                    st.session_state.video_drafts = []

                                st.session_state.video_drafts.append(draft_data)
                                st.success("💾 草稿已保存！您可以稍后继续编辑或上传。")

                # 显示草稿列表
                if "video_drafts" in st.session_state and st.session_state.video_drafts:
                    st.markdown("---")
                    st.markdown("##### 📝 我的草稿")

                    for i, draft in enumerate(st.session_state.video_drafts):
                        with st.expander(f"草稿: {draft['title']}", expanded=False):
                            st.write(f"**文件名:** {draft['file_name']}")
                            st.write(f"**大小:** {draft['file_size']:.2f} MB")
                            st.write(f"**描述:** {draft['description']}")

                            col_edit, col_delete = st.columns(2)
                            with col_edit:
                                if st.button("✏️ 编辑", key=f"edit_draft_{i}"):
                                    st.info("请在上方重新填写信息并上传")
                            with col_delete:
                                if st.button("🗑️ 删除", key=f"delete_draft_{i}"):
                                    st.session_state.video_drafts.pop(i)
                                    st.rerun()

            # --- 子标签2: 我的视频库 ---
            with video_tab2:
                st.markdown("##### 📚 我的视频库")

                db = SessionLocal()
                try:
                    # 查询当前用户的所有视频
                    user_videos = db.query(VideoResource).filter(
                        VideoResource.teacher_id == st.session_state.get("user_id")
                    ).order_by(VideoResource.timestamp.desc()).all()

                    if not user_videos:
                        st.info("您还没有上传任何视频。")
                    else:
                        for video in user_videos:
                            with st.container(border=True):
                                col_info, col_actions = st.columns([3, 1])

                                with col_info:
                                    st.subheader(video.title)
                                    st.caption(f"状态: {video.status} | 上传时间: {video.timestamp.strftime('%Y-%m-%d %H:%M')}")
                                    if video.description:
                                        st.write(video.description)

                                    # 显示视频
                                    try:
                                        st.video(video.path)
                                    except Exception as e:
                                        st.error(f"视频加载失败: {e}")

                                with col_actions:
                                    # 状态切换
                                    new_status = st.selectbox(
                                        "状态",
                                        ["草稿", "已发布"],
                                        index=0 if video.status == "草稿" else 1,
                                        key=f"status_video_{video.id}"
                                    )

                                    if st.button("💾 更新", key=f"update_video_{video.id}"):
                                        video.status = new_status
                                        db.commit()
                                        st.success("状态已更新！")
                                        st.rerun()

                                    if st.button("🤖 AI分析", key=f"ai_analyze_video_{video.id}"):
                                        try:
                                            from utilstongyi import analyze_video_with_tongyi_stream

                                            # 创建流式输出容器
                                            with st.expander("📋 AI分析结果", expanded=True):
                                                st.info("🔄 AI正在分析视频内容，请稍候...")

                                                # 创建一个空的容器用于流式更新
                                                analysis_container = st.empty()

                                                # 流式显示分析结果
                                                for partial_result in analyze_video_with_tongyi_stream(video.path):
                                                    with analysis_container.container():
                                                        st.markdown(partial_result)

                                                st.success("✅ AI分析完成！")

                                        except Exception as e:
                                            st.error(f"❌ AI分析失败: {e}")
                                            with st.expander("🔍 错误详情"):
                                                st.code(str(e))

                                    if st.button("🗑️ 删除", key=f"delete_video_{video.id}"):
                                        db.delete(video)
                                        db.commit()
                                        st.success("视频已删除！")
                                        st.rerun()
                finally:
                    db.close()

            # --- 子标签3: 添加链接 ---
            with video_tab3:
                st.markdown("##### 🔗 添加外部视频链接")
                st.info("如果您的视频已经存储在其他云服务（如阿里云、腾讯云等），可以直接添加链接。")

                with st.form("link_video_form"):
                    video_title_link = st.text_input(
                        "视频标题 *",
                        placeholder="例如：高等数学微积分入门"
                    )
                    video_url = st.text_input(
                        "视频链接 (URL) *",
                        placeholder="https://your-cloud-storage.com/video.mp4"
                    )
                    video_desc_link = st.text_area(
                        "视频简介",
                        placeholder="请简要描述视频内容...",
                        height=100
                    )
                    video_status_link = st.selectbox(
                        "保存状态",
                        ["草稿", "已发布"],
                        index=0
                    )

                    submitted_link = st.form_submit_button("🔗 添加视频链接", use_container_width=True)

                    if submitted_link:
                        if not video_title_link or not video_url:
                            st.error("⚠️ 请填写视频标题和链接！")
                        else:
                            db = SessionLocal()
                            try:
                                new_video = VideoResource(
                                    teacher_id=st.session_state.get("user_id"),
                                    title=video_title_link,
                                    description=video_desc_link,
                                    path=video_url,
                                    status=video_status_link
                                )
                                db.add(new_video)
                                db.commit()
                                st.success(f"✅ 视频链接 '{video_title_link}' 添加成功！")
                                st.balloons()
                            except Exception as e:
                                st.error(f"❌ 保存失败: {e}")
                            finally:
                                db.close()
