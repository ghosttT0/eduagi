# views/resource_management_view.py (统一资源中心版)
import streamlit as st
import pandas as pd
import json
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from database import SessionLocal, TeachingPlan, Exam, User, ExamQuestion, VideoResource


def write_content_to_docx(document, content, level=0):
    """递归处理复杂数据并写入Word文档"""
    if isinstance(content, dict):
        for key, value in content.items():
            # 添加标题
            title = DocumentTitleMap.get(key, key)
            if level == 0:
                heading = document.add_heading(title, level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                document.add_heading(title, level=level)

            # 递归处理内容
            write_content_to_docx(document, value, level + 1)

    elif isinstance(content, list):
        for item in content:
            if isinstance(item, dict) and "question" in item:
                # 特殊处理题目列表
                write_question_to_docx(document, item)
            else:
                write_content_to_docx(document, item, level)

    else:
        # 普通文本内容
        p = document.add_paragraph(str(content))


def write_question_to_docx(document, question_data):
    """将题目数据写入Word文档"""
    # 添加题目编号和类型
    question_type = question_data.get("question_type", "题目")
    document.add_paragraph(f"{question_type}: {question_data.get('question', '')}")

    # 添加选项（如果有）
    if "options" in question_data:
        options = question_data["options"]
        for idx, option in enumerate(options):
            document.add_paragraph(f"{chr(65 + idx)}. {option}")

    # 添加答案（可选）
    if "answer" in question_data:
        document.add_paragraph(f"答案: {question_data['answer']}")

    document.add_paragraph("-" * 50)


# 教案内容标题映射
DocumentTitleMap = {
    "teaching_objectives": "教学目标",
    "teaching_content": "教学内容",
    "key_points": "教学重点",
    "teaching_difficulties": "教学难点",
    "teaching_methods": "教学方法",
    "teaching_process": "教学过程",
    "homework": "课后作业",
    "teaching_reflection": "教学反思"
}


def render():
    """渲染统一的课件与试卷资源管理页面"""
    st.title("📚 教学资源管理中心")
    st.info("在这里，您可以集中管理平台所有的教案和试卷资源。")

    resource_type_tab, exam_type_tab, video_type_tab = st.tabs(["**教案资源**", "**试卷资源**", "**视频资源**"])
    db = SessionLocal()
    try:
        # --- 教案资源选项卡 ---
        with resource_type_tab:
            st.subheader("所有已生成的教案")

            # 从数据库加载所有教案，并关联教师信息
            all_plans_query = db.query(
                TeachingPlan.id,
                TeachingPlan.input_prompt,
                TeachingPlan.output_content,
                TeachingPlan.timestamp,
                User.display_name
            ).join(User, User.id == TeachingPlan.teacher_id).order_by(TeachingPlan.timestamp.desc())

            all_plans = all_plans_query.all()

            if not all_plans:
                st.warning("目前数据库中没有任何教案资源。")
            else:
                # 创建用于展示和筛选的DataFrame
                df = pd.DataFrame(all_plans, columns=['ID', '主题', '内容JSON', '创建时间', '创建教师'])
                df['创建时间'] = pd.to_datetime(df['创建时间']).dt.strftime('%Y-%m-%d %H:%M')

                # 搜索功能
                search_term = st.text_input("🔍 搜索教案主题或内容：", placeholder="例如：卷积神经网络、Python入门...")
                if search_term:
                    df_filtered = df[df['主题'].str.contains(search_term, case=False, na=False)]
                else:
                    df_filtered = df

                # 展示筛选后的数据表格
                st.dataframe(df_filtered[['ID', '创建时间', '创建教师', '主题']], use_container_width=True,
                             hide_index=True)

                st.divider()

                # 预览和导出选中的资源
                if not df_filtered.empty:
                    st.subheader("预览与导出")
                    selected_id = st.selectbox("请从筛选结果中选择一个教案ID进行操作：",
                                               options=df_filtered['ID'].tolist())

                    if selected_id:
                        # 获取选中教案的完整数据
                        selected_plan_series = df[df['ID'] == selected_id].iloc[0]
                        plan_details = json.loads(selected_plan_series['内容JSON'])

                        with st.expander("预览教案详细内容", expanded=False):
                            for key, value in plan_details.items():
                                title = DocumentTitleMap.get(key, key)
                                st.markdown(f"**{title}**")
                                st.markdown(str(value))

                        # Word导出逻辑
                        try:
                            document = Document()

                            # 添加封面标题
                            title = document.add_heading(plan_details.get("title", f"教案ID_{selected_id}"), level=0)
                            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

                            # 添加基本信息
                            document.add_paragraph(f"创建教师: {selected_plan_series['创建教师']}")
                            document.add_paragraph(f"创建时间: {selected_plan_series['创建时间']}")
                            document.add_paragraph("-" * 50)

                            # 写入教案内容
                            write_content_to_docx(document, plan_details)

                            # 保存到内存流
                            file_stream = BytesIO()
                            document.save(file_stream)
                            file_stream.seek(0)

                            st.download_button(
                                label=f"📄 导出ID: {selected_id} 为Word文档",
                                data=file_stream,
                                file_name=f"教案_{selected_id}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            )
                        except Exception as e:
                            st.error(f"创建Word文档时出错: {e}")
                else:
                    st.info("根据您的搜索条件，没有找到匹配的教案。")

        # --- 试卷资源选项卡 ---
        with exam_type_tab:
            st.subheader("所有已发布的试卷")

            # 从数据库加载所有试卷，并关联教师信息
            all_exams_query = db.query(
                Exam.id, Exam.scope, Exam.timestamp, User.display_name
            ).join(User, User.id == Exam.teacher_id).order_by(Exam.timestamp.desc())

            all_exams = all_exams_query.all()

            if not all_exams:
                st.warning("目前数据库中没有任何已发布的试卷。")
            else:
                df_exams = pd.DataFrame(all_exams, columns=['试卷ID', '考察范围', '发布时间', '出题教师'])
                df_exams['发布时间'] = pd.to_datetime(df_exams['发布时间']).dt.strftime('%Y-%m-%d %H:%M')
                st.dataframe(df_exams, use_container_width=True, hide_index=True)

                st.divider()

                # 提供预览功能
                exam_ids = [e[0] for e in all_exams]
                selected_exam_id = st.selectbox("请选择一个试卷ID进行预览：", options=exam_ids)

                if selected_exam_id:
                    questions = db.query(ExamQuestion).filter(ExamQuestion.exam_id == selected_exam_id).all()
                    with st.expander(f"预览试卷ID: {selected_exam_id} 的内容", expanded=True):
                        for i, q in enumerate(questions):
                            st.markdown(f"**第 {i + 1} 题 ({q.question_type})**")
                            st.markdown(q.question_text)
                            if q.question_type == 'multiple_choice':
                                opts = json.loads(q.options)
                                st.radio("选项", opts, key=f"opt_{q.id}", disabled=True)
                            st.success(f"答案: {q.answer}")
                            st.markdown("---")

                    # 导出试卷为Word文档
                    try:
                        document = Document()

                        # 添加封面标题
                        title = document.add_heading(f"试卷: {selected_exam_id}", level=0)
                        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

                        # 添加基本信息
                        exam_info = df_exams[df_exams['试卷ID'] == selected_exam_id].iloc[0]
                        document.add_paragraph(f"考察范围: {exam_info['考察范围']}")
                        document.add_paragraph(f"出题教师: {exam_info['出题教师']}")
                        document.add_paragraph(f"发布时间: {exam_info['发布时间']}")
                        document.add_paragraph("-" * 50)

                        # 写入题目
                        for i, q in enumerate(questions):
                            document.add_heading(f"第 {i + 1} 题 ({q.question_type})", level=2)
                            document.add_paragraph(q.question_text)

                            if q.question_type == 'multiple_choice':
                                opts = json.loads(q.options)
                                for idx, opt in enumerate(opts):
                                    document.add_paragraph(f"{chr(65 + idx)}. {opt}")

                            # 答案单独页
                            if i == 0:
                                document.add_page_break()
                                document.add_heading("答案", level=1)

                            document.add_paragraph(f"第 {i + 1} 题答案: {q.answer}")

                        # 保存到内存流
                        file_stream = BytesIO()
                        document.save(file_stream)
                        file_stream.seek(0)

                        st.download_button(
                            label=f"📄 导出试卷ID: {selected_exam_id} 为Word文档",
                            data=file_stream,
                            file_name=f"试卷_{selected_exam_id}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        )
                    except Exception as e:
                        st.error(f"导出试卷时出错: {e}")

        # --- 视频资源选项卡 ---
        with video_type_tab:
            st.subheader("📹 平台视频资源管理")
            st.info("管理平台所有教师上传的视频资源，支持查看、下载和状态管理。")

            # 统计信息
            total_videos = db.query(VideoResource).count()
            published_videos = db.query(VideoResource).filter(VideoResource.status == "已发布").count()
            draft_videos = db.query(VideoResource).filter(VideoResource.status == "草稿").count()

            col_stat1, col_stat2, col_stat3 = st.columns(3)
            with col_stat1:
                st.metric("📊 总视频数", total_videos)
            with col_stat2:
                st.metric("✅ 已发布", published_videos)
            with col_stat3:
                st.metric("📝 草稿", draft_videos)

            st.markdown("---")

            # 筛选和搜索
            col_filter1, col_filter2, col_filter3 = st.columns(3)

            with col_filter1:
                status_filter = st.selectbox(
                    "📂 状态筛选",
                    ["全部", "已发布", "草稿"]
                )

            with col_filter2:
                teacher_filter = st.selectbox(
                    "👨‍🏫 教师筛选",
                    ["全部教师"] + get_all_video_teachers(db)
                )

            with col_filter3:
                search_video = st.text_input("🔍 搜索视频", placeholder="输入视频标题...")

            # 获取视频列表
            video_query = db.query(VideoResource, User.display_name).join(
                User, User.id == VideoResource.teacher_id
            )

            # 应用筛选条件
            if status_filter != "全部":
                video_query = video_query.filter(VideoResource.status == status_filter)

            if teacher_filter != "全部教师":
                video_query = video_query.filter(User.display_name == teacher_filter)

            if search_video:
                video_query = video_query.filter(
                    (VideoResource.title.contains(search_video)) |
                    (VideoResource.description.contains(search_video))
                )

            videos = video_query.order_by(VideoResource.timestamp.desc()).all()

            if not videos:
                st.warning("📹 没有找到符合条件的视频")
            else:
                st.success(f"📊 找到 {len(videos)} 个视频")

                # 批量操作
                st.markdown("##### 🔧 批量操作")
                col_batch1, col_batch2, col_batch3 = st.columns(3)

                with col_batch1:
                    if st.button("📢 批量发布选中视频", use_container_width=True):
                        st.session_state["batch_publish"] = True

                with col_batch2:
                    if st.button("📝 批量设为草稿", use_container_width=True):
                        st.session_state["batch_draft"] = True

                with col_batch3:
                    if st.button("🗑️ 批量删除选中", use_container_width=True):
                        st.session_state["batch_delete"] = True

                st.markdown("---")

                # 视频列表
                selected_videos = []
                for video, teacher_name in videos:
                    with st.container(border=True):
                        col_check, col_info, col_actions = st.columns([1, 4, 2])

                        with col_check:
                            if st.checkbox("选择", key=f"video_select_{video.id}"):
                                selected_videos.append(video.id)

                        with col_info:
                            # 视频信息
                            st.subheader(f"🎬 {video.title}")
                            time_str = video.timestamp.strftime('%Y-%m-%d %H:%M') if video.timestamp else "时间未知"
                            st.caption(f"👨‍🏫 {teacher_name} | 📅 {time_str} | 📊 {video.status}")

                            if video.description:
                                st.write(f"📝 **简介**: {video.description}")

                            # 视频链接信息
                            if video.path.startswith("http"):
                                st.info(f"🔗 **链接**: {video.path}")
                            else:
                                st.info(f"📁 **文件**: {video.path}")

                            # 显示视频预览
                            if st.button(f"👁️ 预览视频", key=f"preview_{video.id}"):
                                st.session_state[f"show_preview_{video.id}"] = True

                            # 视频预览
                            if st.session_state.get(f"show_preview_{video.id}", False):
                                with st.expander("📺 视频预览", expanded=True):
                                    try:
                                        st.video(video.path)
                                    except Exception as e:
                                        st.error(f"视频加载失败: {e}")

                                    if st.button("❌ 关闭预览", key=f"close_preview_{video.id}"):
                                        st.session_state[f"show_preview_{video.id}"] = False
                                        st.rerun()

                        with col_actions:
                            # 状态管理
                            new_status = st.selectbox(
                                "状态",
                                ["已发布", "草稿"],
                                index=0 if video.status == "已发布" else 1,
                                key=f"status_{video.id}"
                            )

                            if st.button("💾 更新状态", key=f"update_status_{video.id}", use_container_width=True):
                                video.status = new_status
                                db.commit()
                                st.success("状态已更新！")
                                st.rerun()

                            # 下载按钮
                            if video.path.startswith("http"):
                                if st.button("📥 下载视频", key=f"download_{video.id}", use_container_width=True):
                                    download_video_from_url(video.path, video.title)
                            else:
                                if st.button("📁 查看文件", key=f"view_file_{video.id}", use_container_width=True):
                                    st.info(f"文件路径: {video.path}")

                            # 删除按钮
                            if st.button("🗑️ 删除", key=f"delete_video_{video.id}", use_container_width=True):
                                if st.session_state.get(f"confirm_delete_video_{video.id}", False):
                                    db.delete(video)
                                    db.commit()
                                    st.success("视频已删除！")
                                    st.rerun()
                                else:
                                    st.session_state[f"confirm_delete_video_{video.id}"] = True
                                    st.warning("再次点击确认删除")

                # 处理批量操作
                if st.session_state.get("batch_publish", False) and selected_videos:
                    for video_id in selected_videos:
                        video = db.query(VideoResource).filter(VideoResource.id == video_id).first()
                        if video:
                            video.status = "已发布"
                    db.commit()
                    st.success(f"已批量发布 {len(selected_videos)} 个视频！")
                    st.session_state["batch_publish"] = False
                    st.rerun()

                if st.session_state.get("batch_draft", False) and selected_videos:
                    for video_id in selected_videos:
                        video = db.query(VideoResource).filter(VideoResource.id == video_id).first()
                        if video:
                            video.status = "草稿"
                    db.commit()
                    st.success(f"已批量设为草稿 {len(selected_videos)} 个视频！")
                    st.session_state["batch_draft"] = False
                    st.rerun()

                if st.session_state.get("batch_delete", False) and selected_videos:
                    for video_id in selected_videos:
                        video = db.query(VideoResource).filter(VideoResource.id == video_id).first()
                        if video:
                            db.delete(video)
                    db.commit()
                    st.success(f"已批量删除 {len(selected_videos)} 个视频！")
                    st.session_state["batch_delete"] = False
                    st.rerun()

    finally:
        db.close()


def get_all_video_teachers(db):
    """获取所有上传过视频的教师列表"""
    try:
        teachers = db.query(User.display_name).join(
            VideoResource, User.id == VideoResource.teacher_id
        ).distinct().all()
        return [teacher.display_name for teacher in teachers]
    except:
        return []


def download_video_from_url(video_url, video_title):
    """从URL下载视频"""
    try:
        import requests
        import os
        from urllib.parse import urlparse

        # 解析URL获取文件扩展名
        parsed_url = urlparse(video_url)
        file_extension = os.path.splitext(parsed_url.path)[1] or '.mp4'

        # 生成安全的文件名
        safe_title = "".join(c for c in video_title if c.isalnum() or c in (' ', '-', '_')).rstrip()
        filename = f"{safe_title}{file_extension}"

        st.info(f"🔄 正在准备下载: {filename}")
        st.info(f"🔗 视频链接: {video_url}")
        st.warning("⚠️ 请手动复制链接到下载工具中下载，或右键另存为")

        # 提供下载链接
        st.markdown(f"[📥 点击下载 {filename}]({video_url})")

    except Exception as e:
        st.error(f"下载准备失败: {e}")
        st.info(f"🔗 原始链接: {video_url}")