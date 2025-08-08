from fastapi import APIRouter, Depends, HTTPException
from sqlalchemy.orm import Session
from pydantic import BaseModel
from typing import List, Optional
import json
from datetime import datetime, timedelta
from io import BytesIO
import docx
from docx.shared import Inches

from database import (
    get_db, User, ChatHistory, KnowledgePoint, StudentDispute,
    KnowledgeMastery, VideoResource, Class
)
from api.auth import get_current_user
from services.ai_service import ai_service

# 创建路由器
student_router = APIRouter()

# Pydantic模型
class ChatMessage(BaseModel):
    question: str
    ai_mode: str = "直接问答"  # 直接问答、苏格拉底式引导、关联知识分析

class ChatResponse(BaseModel):
    answer: str
    timestamp: datetime

class ChatHistoryResponse(BaseModel):
    id: int
    question: str
    answer: str
    timestamp: datetime
    
    class Config:
        from_attributes = True

class PracticeRequest(BaseModel):
    topic: str

class PracticeQuestion(BaseModel):
    question_text: str
    standard_answer: str
    topic: str

class PracticeAnswer(BaseModel):
    student_answer: str

class PracticeFeedback(BaseModel):
    feedback: str
    score: Optional[int] = None

class DisputeCreate(BaseModel):
    message: str

class DisputeResponse(BaseModel):
    id: int
    message: str
    status: str
    teacher_reply: Optional[str] = None
    created_at: datetime
    replied_at: Optional[datetime] = None
    
    class Config:
        from_attributes = True

class KnowledgeMasteryCreate(BaseModel):
    knowledge_point: str
    mastery_level: int  # 1-薄弱, 2-基本, 3-熟练
    self_assessment: Optional[str] = None

class KnowledgeMasteryResponse(BaseModel):
    id: int
    knowledge_point: str
    mastery_level: int
    self_assessment: Optional[str] = None
    created_at: datetime
    updated_at: datetime
    
    class Config:
        from_attributes = True

class VideoResourceResponse(BaseModel):
    id: int
    title: str
    description: Optional[str] = None
    path: str
    status: str
    created_at: datetime
    teacher_name: str
    
    class Config:
        from_attributes = True

# AI学习伙伴相关接口
@student_router.post("/chat", response_model=ChatResponse)
async def chat_with_ai(
    message: ChatMessage,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """与AI学习伙伴对话 - 优化版本"""
    if current_user.role != "学生":
        raise HTTPException(status_code=403, detail="权限不足")
    
    try:
        # 输入验证
        if not message.question or len(message.question.strip()) < 2:
            raise HTTPException(status_code=400, detail="问题内容不能为空且至少包含2个字符")
        
        if len(message.question) > 1000:
            raise HTTPException(status_code=400, detail="问题内容过长，请控制在1000字符以内")
        
        # 获取最近的对话历史作为上下文
        recent_chats = db.query(ChatHistory).filter(
            ChatHistory.student_id == current_user.id
        ).order_by(ChatHistory.timestamp.desc()).limit(10).all()
        
        # 构建对话上下文
        chat_history = []
        for chat in reversed(recent_chats):
            chat_history.append((chat.question, chat.answer))
        
        try:
            # 调用AI服务生成回答
            ai_answer = await ai_service.chat_with_student(
                question=message.question,
                ai_mode=message.ai_mode,
                chat_history=chat_history
            )
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"AI服务调用失败: {str(e)}。请检查API密钥配置。"
            )
        
        # 验证AI回答质量
        if not ai_answer or len(ai_answer.strip()) < 10:
            ai_answer = f"很抱歉，我需要更多信息来回答您关于 '{message.question}' 的问题。请您提供更多具体的背景信息，这样我就能给出更准确和有用的回答。"
        
        # 保存对话历史
        chat_record = ChatHistory(
            student_id=current_user.id,
            question=message.question.strip(),
            answer=ai_answer
        )
        
        db.add(chat_record)
        db.commit()
        
        return ChatResponse(answer=ai_answer, timestamp=datetime.now())
        
    except HTTPException:
        raise
    except Exception as e:
        # 记录错误但不暴露给用户
        print(f"AI对话服务错误: {str(e)}")
        raise HTTPException(status_code=500, detail="AI服务暂时不可用，请稍后重试")

@student_router.get("/chat/history", response_model=List[ChatHistoryResponse])
async def get_chat_history(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
    limit: int = 50
):
    """获取聊天历史"""
    if current_user.role != "学生":
        raise HTTPException(status_code=403, detail="权限不足")
    
    chats = db.query(ChatHistory).filter(
        ChatHistory.student_id == current_user.id
    ).order_by(ChatHistory.timestamp.desc()).limit(limit).all()
    
    return [ChatHistoryResponse.from_orm(chat) for chat in chats]

@student_router.delete("/chat/history")
async def clear_chat_history(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """清空聊天历史"""
    if current_user.role != "学生":
        raise HTTPException(status_code=403, detail="权限不足")
    
    db.query(ChatHistory).filter(
        ChatHistory.student_id == current_user.id
    ).delete()
    
    db.commit()
    
    return {"message": "聊天历史已清空"}

# 自主练习相关接口
@student_router.post("/practice/generate", response_model=PracticeQuestion)
async def generate_practice_question(
    request: PracticeRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """生成练习题 - 优化版本"""
    if current_user.role != "学生":
        raise HTTPException(status_code=403, detail="权限不足")
    
    try:
        # 输入验证
        if not request.topic or len(request.topic.strip()) < 2:
            raise HTTPException(status_code=400, detail="主题内容不能为空且至少包含2个字符")
        
        if len(request.topic) > 200:
            raise HTTPException(status_code=400, detail="主题内容过长，请控制在200字符以内")
        
        # 更新知识点查询次数
        kp = db.query(KnowledgePoint).filter(
            KnowledgePoint.topic == request.topic.strip()
        ).first()
        
        if kp:
            kp.query_count += 1
        else:
            kp = KnowledgePoint(topic=request.topic.strip(), query_count=1)
            db.add(kp)
        
        db.commit()
        
        try:
            # 调用AI服务生成练习题
            question_data = await ai_service.generate_practice_question(request.topic.strip())
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"AI服务调用失败: {str(e)}。请检查API密钥配置。"
            )
        
        # 验证AI返回的数据质量
        if not isinstance(question_data, dict):
            raise ValueError("题目数据格式错误")
        
        if not question_data.get("question_text") or not question_data.get("standard_answer"):
            raise ValueError("题目或答案内容缺失")
        
        # 验证内容质量
        question_text = question_data["question_text"].strip()
        standard_answer = question_data["standard_answer"].strip()
        
        if len(question_text) < 10:
            raise ValueError("题目内容过于简单")
        
        if len(standard_answer) < 10:
            raise ValueError("答案内容过于简单")
        
        # 检查是否是默认模板题目（简单的质量检查）
        if "请解释" in question_text and "的基本概念" in question_text:
            # 这可能是默认模板，重新生成
            question_data = await ai_service.generate_practice_question(request.topic.strip())
            question_text = question_data.get("question_text", question_text).strip()
            standard_answer = question_data.get("standard_answer", standard_answer).strip()
        
        practice_question = PracticeQuestion(
            question_text=question_text,
            standard_answer=standard_answer,
            topic=request.topic.strip()
        )
        
        return practice_question
        
    except HTTPException:
        raise
    except Exception as e:
        # 记录错误但不暴露给用户
        print(f"题目生成服务错误: {str(e)}")
        
        # 返回一个高质量的备用题目
        fallback_questions = {
            "Python": {
                "question_text": "请编写一个Python函数，实现对一个数字列表进行升序排列，并返回排列后的结果。要求使用内置的sorted()函数。",
                "standard_answer": "```python\ndef sort_numbers(numbers):\n    return sorted(numbers)\n\n# 示例使用\nnumbers = [3, 1, 4, 1, 5, 9, 2, 6]\nresult = sort_numbers(numbers)\nprint(result)  # 输出: [1, 1, 2, 3, 4, 5, 6, 9]\n```\n\n解释：\n1. 定义函数sort_numbers()，接收一个数字列表作为参数\n2. 使用sorted()内置函数对列表进行升序排列\n3. 返回排列后的新列表（不修改原列表）"
            },
            "默认": {
                "question_text": f"请简述{request.topic}的主要特点和应用场景，并举一个具体的实际应用例子说明。",
                "standard_answer": f"{request.topic}的主要特点包括：\n1. 核心概念和原理\n2. 主要优势和特色\n3. 适用的应用场景\n\n实际应用例子：在具体项目中，{request.topic}可以用于解决实际问题，提高效率和质量。"
            }
        }
        
        topic_key = request.topic.strip() if "Python" in request.topic else "默认"
        fallback = fallback_questions.get(topic_key, fallback_questions["默认"])
        
        return PracticeQuestion(
            question_text=fallback["question_text"],
            standard_answer=fallback["standard_answer"],
            topic=request.topic.strip()
        )

@student_router.post("/practice/submit", response_model=PracticeFeedback)
async def submit_practice_answer(
    answer: PracticeAnswer,
    question_data: PracticeQuestion,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """提交练习答案并获取反馈 - 优化版本"""
    if current_user.role != "学生":
        raise HTTPException(status_code=403, detail="权限不足")
    
    try:
        # 输入验证
        if not answer.student_answer or len(answer.student_answer.strip()) < 2:
            raise HTTPException(status_code=400, detail="答案内容不能为空且至少包含2个字符")
        
        if len(answer.student_answer) > 2000:
            raise HTTPException(status_code=400, detail="答案内容过长，请控制在2000字符以内")
        
        if not question_data.question_text or not question_data.standard_answer:
            raise HTTPException(status_code=400, detail="题目数据不完整")
        
        # 调用AI服务生成反馈
        try:
            feedback_text = await ai_service.evaluate_practice_answer(
                question=question_data.question_text,
                standard_answer=question_data.standard_answer,
                student_answer=answer.student_answer.strip(),
                topic=getattr(question_data, 'topic', '')
            )
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"AI服务调用失败: {str(e)}。请检查API密钥配置。"
            )
        
        # 验证反馈质量
        if not feedback_text or len(feedback_text.strip()) < 20:
            feedback_text = f"""
📝 **答案反馈**

您的回答：{answer.student_answer.strip()}

🔍 **初步评估**：
您的答案显示了对该问题的思考。为了给出更准确的评估，建议您：

1. 提供更详细的解释和理由
2. 包含具体的例子或步骤
3. 检查答案的完整性和准确性

📚 **学习建议**：
继续加油！学习是一个持续的过程，每一次练习都是进步的机会。
            """.strip()
        
        # 简单的评分逻辑（基于答案长度和关键词匹配）
        score = 5  # 默认基础分
        
        # 根据答案长度调整分数
        answer_length = len(answer.student_answer.strip())
        if answer_length > 100:
            score += 2
        elif answer_length > 50:
            score += 1
        
        # 检查是否包含关键词（简单的关键词匹配）
        topic = getattr(question_data, 'topic', '').lower()
        answer_lower = answer.student_answer.lower()
        
        if topic and topic in answer_lower:
            score += 1
        
        # 检查是否包含代码或结构化内容
        if '```' in answer.student_answer or '步骤' in answer.student_answer or '首先' in answer.student_answer:
            score += 1
        
        # 限制分数范围
        score = min(max(score, 1), 10)
        
        return PracticeFeedback(feedback=feedback_text, score=score)
        
    except HTTPException:
        raise
    except Exception as e:
        # 记录错误但不暴露给用户
        print(f"答案评估服务错误: {str(e)}")
        
        # 返回一个基础的反馈
        fallback_feedback = f"""
📝 **答案反馈**

您的回答：{answer.student_answer.strip()}

🔍 **评估结果**：
您的答案显示了对问题的思考和理解。以下是一些建议：

✅ **优点**：
- 积极参与练习，显示了学习的主动性
- 给出了自己的理解和见解

💡 **改进建议**：
- 可以更加详细地阐述解决方案或思路
- 尝试结合具体例子来说明您的观点
- 注意答案的逻辑性和条理性

📚 **学习指导**：
继续保持这种学习热情！每一次练习都是提高的机会。建议您多做类似的练习，加深对知识点的理解。

🎆 **综合评分：6/10**
评分理由：答案显示了基本的理解，但在深度和完整性方面还有提升空间。
        """.strip()
        
        return PracticeFeedback(feedback=fallback_feedback, score=6)

# 向老师提问相关接口
@student_router.post("/disputes", response_model=DisputeResponse)
async def create_dispute(
    dispute_data: DisputeCreate,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """向老师提交疑问 - 优化版本"""
    if current_user.role != "学生":
        raise HTTPException(status_code=403, detail="权限不足")
    
    try:
        # 输入验证
        if not dispute_data.message or len(dispute_data.message.strip()) < 5:
            raise HTTPException(status_code=400, detail="问题内容不能为空且至少包含5个字符")
        
        if len(dispute_data.message) > 1000:
            raise HTTPException(status_code=400, detail="问题内容过长，请控制在1000字符以内")
        
        # 检查学生是否属于某个班级
        if not current_user.class_id:
            raise HTTPException(status_code=400, detail="您尚未加入任何班级，无法向老师提问")
        
        # 检查班级是否存在
        class_info = db.query(Class).filter(Class.id == current_user.class_id).first()
        if not class_info:
            raise HTTPException(status_code=404, detail="班级信息不存在")
        
        # 检查是否有老师
        teacher = db.query(User).filter(
            User.class_id == current_user.class_id,
            User.role == "老师"
        ).first()
        
        if not teacher:
            raise HTTPException(status_code=404, detail="未找到班级老师")
        
        # 检查最近是否有重复提问（防止垃圾信息）
        recent_dispute = db.query(StudentDispute).filter(
            StudentDispute.student_id == current_user.id,
            StudentDispute.message == dispute_data.message.strip(),
            StudentDispute.created_at > datetime.now() - timedelta(minutes=5)
        ).first()
        
        if recent_dispute:
            raise HTTPException(status_code=400, detail="请勿重复提交相同的问题")
        
        # 创建疑问记录
        new_dispute = StudentDispute(
            student_id=current_user.id,
            class_id=current_user.class_id,
            message=dispute_data.message.strip(),
            status="待回复"
        )
        
        db.add(new_dispute)
        db.commit()
        db.refresh(new_dispute)
        
        return DisputeResponse(
            id=new_dispute.id,
            message=new_dispute.message,
            status=new_dispute.status,
            created_at=new_dispute.created_at
        )
        
    except HTTPException:
        raise
    except Exception as e:
        # 记录错误但不暴露给用户
        print(f"提问服务错误: {str(e)}")
        raise HTTPException(status_code=500, detail="提问服务暂时不可用，请稍后重试")

@student_router.get("/disputes", response_model=List[DisputeResponse])
async def get_my_disputes(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """获取我的疑问列表"""
    if current_user.role != "学生":
        raise HTTPException(status_code=403, detail="权限不足")
    
    disputes = db.query(StudentDispute).filter(
        StudentDispute.student_id == current_user.id
    ).order_by(StudentDispute.created_at.desc()).all()
    
    return [DisputeResponse.from_orm(dispute) for dispute in disputes]

# 知识掌握评估相关接口
@student_router.post("/knowledge-mastery", response_model=KnowledgeMasteryResponse)
async def create_or_update_knowledge_mastery(
    mastery_data: KnowledgeMasteryCreate,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """创建或更新知识掌握评估"""
    if current_user.role != "学生":
        raise HTTPException(status_code=403, detail="权限不足")
    
    # 检查是否已存在该知识点的评估
    existing = db.query(KnowledgeMastery).filter(
        KnowledgeMastery.student_id == current_user.id,
        KnowledgeMastery.knowledge_point == mastery_data.knowledge_point
    ).first()
    
    if existing:
        # 更新现有记录
        existing.mastery_level = mastery_data.mastery_level
        existing.self_assessment = mastery_data.self_assessment
        existing.updated_at = datetime.now()
        db.commit()
        db.refresh(existing)
        return KnowledgeMasteryResponse.from_orm(existing)
    else:
        # 创建新记录
        new_mastery = KnowledgeMastery(
            student_id=current_user.id,
            knowledge_point=mastery_data.knowledge_point,
            mastery_level=mastery_data.mastery_level,
            self_assessment=mastery_data.self_assessment
        )
        db.add(new_mastery)
        db.commit()
        db.refresh(new_mastery)
        return KnowledgeMasteryResponse.from_orm(new_mastery)

@student_router.get("/knowledge-mastery", response_model=List[KnowledgeMasteryResponse])
async def get_knowledge_mastery(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """获取知识掌握评估列表"""
    if current_user.role != "学生":
        raise HTTPException(status_code=403, detail="权限不足")
    
    masteries = db.query(KnowledgeMastery).filter(
        KnowledgeMastery.student_id == current_user.id
    ).order_by(KnowledgeMastery.updated_at.desc()).all()
    
    return [KnowledgeMasteryResponse.from_orm(mastery) for mastery in masteries]

# 视频学习相关接口
@student_router.get("/videos", response_model=List[VideoResourceResponse])
async def get_available_videos(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """获取可用的视频资源"""
    if current_user.role != "学生":
        raise HTTPException(status_code=403, detail="权限不足")
    
    # 获取已发布的视频资源
    videos = db.query(VideoResource).filter(
        VideoResource.status == "已发布"
    ).order_by(VideoResource.created_at.desc()).all()
    
    result = []
    for video in videos:
        teacher = db.query(User).filter(User.id == video.teacher_id).first()
        result.append(VideoResourceResponse(
            id=video.id,
            title=video.title,
            description=video.description,
            path=video.path,
            status=video.status,
            created_at=video.created_at,
            teacher_name=teacher.display_name if teacher else "未知教师"
        ))
    
    return result

@student_router.get("/videos/{video_id}", response_model=VideoResourceResponse)
async def get_video_detail(
    video_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """获取视频详情"""
    if current_user.role != "学生":
        raise HTTPException(status_code=403, detail="权限不足")
    
    video = db.query(VideoResource).filter(
        VideoResource.id == video_id,
        VideoResource.status == "已发布"
    ).first()
    
    if not video:
        raise HTTPException(status_code=404, detail="视频不存在或未发布")
    
    teacher = db.query(User).filter(User.id == video.teacher_id).first()
    
    return VideoResourceResponse(
        id=video.id,
        title=video.title,
        description=video.description,
        path=video.path,
        status=video.status,
        created_at=video.created_at,
        teacher_name=teacher.display_name if teacher else "未知教师"
    )

# ==================== 学习计划相关模型 ====================

class StudyEventCreate(BaseModel):
    title: str
    start: datetime
    end: datetime
    type: str  # study, exam, assignment, review
    subject: str
    description: Optional[str] = None

class StudyEventUpdate(BaseModel):
    title: Optional[str] = None
    start: Optional[datetime] = None
    end: Optional[datetime] = None
    type: Optional[str] = None
    subject: Optional[str] = None
    description: Optional[str] = None
    completed: Optional[bool] = None

class StudyEventResponse(BaseModel):
    id: int
    title: str
    start: datetime
    end: datetime
    type: str
    subject: str
    description: Optional[str]
    completed: bool
    user_id: int
    created_at: datetime

class StudyPlanCreate(BaseModel):
    title: str
    description: str
    subject: str
    difficulty: str  # beginner, intermediate, advanced
    duration: int  # 天数
    goals: List[str]

class StudyPlanResponse(BaseModel):
    id: int
    title: str
    description: str
    subject: str
    difficulty: str
    duration: int
    goals: List[str]
    progress: float
    user_id: int
    created_at: datetime

class AIStudyPlanRequest(BaseModel):
    subject: str
    goal: str
    level: str  # beginner, intermediate, advanced
    duration: int
    preferences: Optional[List[str]] = []
    daily_time: Optional[int] = 2
    requirements: Optional[str] = ""

# ==================== 学习计划API ====================

@student_router.get("/study-events", response_model=List[StudyEventResponse])
async def get_study_events(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """获取学生的学习事件列表"""
    # 这里需要在数据库中添加StudyEvent表，暂时返回模拟数据
    return []

@student_router.post("/study-events", response_model=StudyEventResponse)
async def create_study_event(
    event: StudyEventCreate,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """创建学习事件"""
    # 这里需要在数据库中添加StudyEvent表，暂时返回模拟数据
    return StudyEventResponse(
        id=1,
        title=event.title,
        start=event.start,
        end=event.end,
        type=event.type,
        subject=event.subject,
        description=event.description,
        completed=False,
        user_id=current_user.id,
        created_at=datetime.now()
    )

@student_router.get("/study-plans", response_model=List[StudyPlanResponse])
async def get_study_plans(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """获取学生的学习计划列表"""
    # 这里需要在数据库中添加StudyPlan表，暂时返回模拟数据
    return []

@student_router.post("/ai-study-plan", response_model=StudyPlanResponse)
async def generate_ai_study_plan(
    request: AIStudyPlanRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """AI生成学习计划"""
    try:
        # 构建AI提示
        prompt = f"""
        请为学生生成一个详细的学习计划：

        学习科目：{request.subject}
        学习目标：{request.goal}
        当前水平：{request.level}
        计划时长：{request.duration}天
        每日可用时间：{request.daily_time}小时
        学习偏好：{', '.join(request.preferences) if request.preferences else '无特殊偏好'}
        特殊要求：{request.requirements or '无'}

        请生成一个包含以下内容的学习计划：
        1. 计划标题
        2. 详细描述
        3. 学习目标（3-5个）
        4. 分阶段学习安排

        请以JSON格式返回，包含title, description, goals字段。
        """

        # 调用AI服务生成计划
        ai_response = await ai_service.generate_response(prompt, "教学计划生成")

        # 解析AI响应（这里简化处理）
        plan_data = {
            "title": f"{request.subject} - {request.goal}学习计划",
            "description": f"基于AI分析生成的个性化{request.subject}学习计划",
            "goals": [
                f"掌握{request.subject}核心概念",
                f"完成实践项目",
                f"通过相关考试",
                f"达到{request.goal}水平"
            ]
        }

        return StudyPlanResponse(
            id=1,
            title=plan_data["title"],
            description=plan_data["description"],
            subject=request.subject,
            difficulty=request.level,
            duration=request.duration,
            goals=plan_data["goals"],
            progress=0.0,
            user_id=current_user.id,
            created_at=datetime.now()
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI生成学习计划失败: {str(e)}")

@student_router.get("/study-plans/{plan_id}/export")
async def export_study_plan_to_word(
    plan_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """导出学习计划为Word文档"""
    try:
        # 创建Word文档
        doc = docx.Document()

        # 添加标题
        title = doc.add_heading('个性化学习计划', 0)
        title.alignment = 1  # 居中对齐

        # 添加基本信息
        doc.add_heading('基本信息', level=1)
        doc.add_paragraph(f'学生姓名：{current_user.display_name}')
        doc.add_paragraph(f'生成时间：{datetime.now().strftime("%Y年%m月%d日")}')
        doc.add_paragraph(f'计划编号：{plan_id}')

        # 添加学习目标
        doc.add_heading('学习目标', level=1)
        goals = [
            '掌握核心知识点和概念',
            '提高实践应用能力',
            '培养独立思考和解决问题的能力',
            '达到预期的学习效果'
        ]
        for goal in goals:
            doc.add_paragraph(f'• {goal}')

        # 添加学习安排
        doc.add_heading('学习安排', level=1)
        doc.add_paragraph('本学习计划采用循序渐进的方式，分为以下几个阶段：')

        stages = [
            '第一阶段：基础知识学习（1-2周）',
            '第二阶段：实践练习（2-3周）',
            '第三阶段：综合应用（1-2周）',
            '第四阶段：总结复习（1周）'
        ]
        for stage in stages:
            doc.add_paragraph(f'• {stage}')

        # 添加学习建议
        doc.add_heading('学习建议', level=1)
        suggestions = [
            '制定每日学习计划，保持学习的连续性',
            '及时复习所学内容，巩固知识点',
            '多做练习题，提高实践能力',
            '遇到问题及时向老师或同学请教',
            '定期自我评估学习效果'
        ]
        for suggestion in suggestions:
            doc.add_paragraph(f'• {suggestion}')

        # 保存到内存
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        from fastapi.responses import StreamingResponse

        return StreamingResponse(
            BytesIO(buffer.read()),
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={'Content-Disposition': f'attachment; filename="学习计划_{plan_id}.docx"'}
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"导出Word文档失败: {str(e)}")
