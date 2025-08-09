from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from pydantic import BaseModel
from typing import List, Optional, Dict, Any
import json
from datetime import datetime
from io import BytesIO

import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from database import (
    get_db, User, TeachingPlan, MindMap, VideoResource,
    StudentDispute, KnowledgePoint
)
from api.auth import get_current_user
from services.ai_service import ai_service

# 创建路由器
teacher_router = APIRouter()

# Pydantic模型
class TeachingPlanCreate(BaseModel):
    course_name: str
    chapter: str
    topic: Optional[str] = None
    class_hours: int = 2
    teaching_time: int = 90

class TeachingPlanResponse(BaseModel):
    id: int
    teacher_id: int
    input_prompt: str
    output_content: str
    created_at: datetime
    
    class Config:
        from_attributes = True

class MindMapCreate(BaseModel):
    title: str
    topic: str
    description: Optional[str] = None
    is_public: bool = False

class MindMapResponse(BaseModel):
    id: int
    user_id: int
    title: str
    topic: str
    data: str
    description: Optional[str] = None
    is_public: bool
    created_at: datetime
    
    class Config:
        from_attributes = True

class VideoResourceCreate(BaseModel):
    title: str
    description: Optional[str] = None
    path: str
    status: str = "草稿"

class VideoResourceResponse(BaseModel):
    id: int
    teacher_id: int
    title: str
    description: Optional[str] = None
    path: str
    status: str
    created_at: datetime
    
    class Config:
        from_attributes = True

class StudentDisputeResponse(BaseModel):
    id: int
    student_id: int
    student_name: str
    class_id: int
    question_id: Optional[int] = None
    message: str
    status: str
    teacher_reply: Optional[str] = None
    created_at: datetime
    replied_at: Optional[datetime] = None
    
    class Config:
        from_attributes = True

class ExamGenerateRequest(BaseModel):
    exam_scope: str
    num_mcq: int = 5
    num_saq: int = 3
    num_code: int = 1

# 智能教学设计相关接口
@teacher_router.post("/teaching-plans", response_model=TeachingPlanResponse)
async def create_teaching_plan(
    plan_data: TeachingPlanCreate,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """创建教学计划"""
    if current_user.role != "教师":
        raise HTTPException(status_code=403, detail="权限不足")
    
    # 调用AI服务生成教学计划
    topic_text = plan_data.topic if plan_data.topic else f"{plan_data.course_name} - {plan_data.chapter} 整体大纲"

    try:
        ai_response = await ai_service.generate_teaching_plan(
            course_name=plan_data.course_name,
            chapter=plan_data.chapter,
            topic=plan_data.topic,
            class_hours=plan_data.class_hours,
            teaching_time=plan_data.teaching_time
        )
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"AI服务调用失败: {str(e)}。请检查API密钥配置。"
        )
    
    # 保存到数据库
    new_plan = TeachingPlan(
        teacher_id=current_user.id,
        input_prompt=topic_text,
        output_content=json.dumps(ai_response, ensure_ascii=False)
    )
    
    db.add(new_plan)
    db.commit()
    db.refresh(new_plan)
    
    return TeachingPlanResponse.from_orm(new_plan)

@teacher_router.get("/teaching-plans", response_model=List[TeachingPlanResponse])
async def get_teaching_plans(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """获取教师的教学计划列表"""
    if current_user.role != "教师":
        raise HTTPException(status_code=403, detail="权限不足")
    
    plans = db.query(TeachingPlan).filter(
        TeachingPlan.teacher_id == current_user.id
    ).order_by(TeachingPlan.created_at.desc()).all()
    
    return [TeachingPlanResponse.from_orm(plan) for plan in plans]

@teacher_router.get("/teaching-plans/{plan_id}", response_model=TeachingPlanResponse)
async def get_teaching_plan(
    plan_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """获取特定教学计划"""
    if current_user.role != "教师":
        raise HTTPException(status_code=403, detail="权限不足")
    
    plan = db.query(TeachingPlan).filter(
        TeachingPlan.id == plan_id,
        TeachingPlan.teacher_id == current_user.id
    ).first()
    
    if not plan:
        raise HTTPException(status_code=404, detail="教学计划不存在")
    
    return TeachingPlanResponse.from_orm(plan)

# 知识图谱相关接口
@teacher_router.post("/mindmaps", response_model=MindMapResponse)
async def create_mindmap(
    mindmap_data: MindMapCreate,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """创建思维导图"""
    if current_user.role != "教师":
        raise HTTPException(status_code=403, detail="权限不足")
    
    try:
        # 调用AI服务生成思维导图
        mindmap_json = await ai_service.generate_mindmap(
            topic=mindmap_data.topic,
            description=mindmap_data.description
        )
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"AI服务调用失败: {str(e)}。请检查API密钥配置。"
        )
    
    new_mindmap = MindMap(
        user_id=current_user.id,
        title=mindmap_data.title,
        topic=mindmap_data.topic,
        data=json.dumps(mindmap_json, ensure_ascii=False),
        description=mindmap_data.description,
        is_public=mindmap_data.is_public
    )
    
    db.add(new_mindmap)
    db.commit()
    db.refresh(new_mindmap)
    
    return MindMapResponse.from_orm(new_mindmap)

@teacher_router.get("/mindmaps", response_model=List[MindMapResponse])
async def get_mindmaps(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """获取教师的思维导图列表"""
    if current_user.role != "教师":
        raise HTTPException(status_code=403, detail="权限不足")
    
    mindmaps = db.query(MindMap).filter(
        MindMap.user_id == current_user.id
    ).order_by(MindMap.created_at.desc()).all()
    
    return [MindMapResponse.from_orm(mindmap) for mindmap in mindmaps]

# 智能出题接口
@teacher_router.post("/generate-exam")
async def generate_exam(
    exam_data: ExamGenerateRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """智能生成试卷"""
    if current_user.role != "教师":
        raise HTTPException(status_code=403, detail="权限不足")
    
    try:
        # 调用AI服务生成试卷
        exam_questions = await ai_service.generate_exam_questions(
            exam_scope=exam_data.exam_scope,
            num_mcq=exam_data.num_mcq,
            num_saq=exam_data.num_saq,
            num_code=exam_data.num_code
        )
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"AI服务调用失败: {str(e)}。请检查API密钥配置。"
        )
    
    return {
        "message": "试卷生成成功",
        "exam_data": exam_questions
    }


# ============== 教学计划 Word 导出（高级版） ==============
@teacher_router.get("/teaching-plans/{plan_id}/export")
async def export_teaching_plan_to_word(
    plan_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """导出教学计划为高级格式的Word文档（包含统一样式与表格）。"""
    if current_user.role != "教师":
        raise HTTPException(status_code=403, detail="权限不足")

    # 获取教学计划
    plan = db.query(TeachingPlan).filter(
        TeachingPlan.id == plan_id,
        TeachingPlan.teacher_id == current_user.id
    ).first()

    if not plan:
        raise HTTPException(status_code=404, detail="教学计划不存在")

    # 解析教学计划JSON内容（健壮解析：处理双重JSON、代码块围栏等情况）
    def robust_parse(content_text: str) -> Dict[str, Any]:
        if not content_text:
            return {}
        try:
            data = json.loads(content_text)
            # 如果第一层解析得到的是字符串，再次尝试解析
            if isinstance(data, str):
                text = data.strip()
                # 去除可能的```json或```围栏
                if text.startswith("```json") and text.endswith("```"):
                    text = text[7:-3].strip()
                elif text.startswith("```") and text.endswith("```"):
                    text = text[3:-3].strip()
                try:
                    return json.loads(text)
                except Exception:
                    return {"教学内容": text}
            if isinstance(data, dict):
                return data
            return {}
        except Exception:
            # 去除可能的```json或```围栏后重试
            text = content_text.strip()
            if text.startswith("```json") and text.endswith("```"):
                text = text[7:-3].strip()
            elif text.startswith("```") and text.endswith("```"):
                text = text[3:-3].strip()
            try:
                return json.loads(text)
            except Exception:
                return {"教学内容": text}

    plan_details = robust_parse(plan.output_content)

    # 兼容老键名（英文）与新键名（中文）
    key_aliases = {
        "教学内容": ["教学内容", "teaching_content"],
        "教学目标": ["教学目标", "teaching_objectives"],
        "教学重点": ["教学重点", "key_points"],
        "教学难点": ["教学难点", "teaching_difficulties"],
        "教学设计": ["教学设计", "teaching_design"],
        "教学反思与总结": ["教学反思与总结", "teaching_reflection", "reflection"],
    }

    def get_value(aliases: List[str]) -> str:
        for k in aliases:
            v = plan_details.get(k)
            if v:
                return str(v)
        return ""

    # 创建文档与统一样式
    doc = docx.Document()

    # 设置全局字体（中文：宋体，英文：Calibri）
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)
    # 设置中文字体（东亚字体）
    try:
        normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    except Exception:
        pass

    # 标题
    title_text = f"《{plan.input_prompt}》教案"
    title_para = doc.add_heading(title_text, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 基本信息表格
    info_table = doc.add_table(rows=3, cols=4)
    info_table.style = 'Light Grid'
    info_table.autofit = True
    cells = info_table.rows
    cells[0].cells[0].text = "教师姓名"
    cells[0].cells[1].text = current_user.display_name if hasattr(current_user, 'display_name') and current_user.display_name else "—"
    cells[0].cells[2].text = "生成日期"
    cells[0].cells[3].text = plan.created_at.strftime("%Y-%m-%d %H:%M") if hasattr(plan, 'created_at') else datetime.now().strftime("%Y-%m-%d %H:%M")

    cells[1].cells[0].text = "教案编号"
    cells[1].cells[1].text = str(plan.id)
    cells[1].cells[2].text = "主题/章节"
    cells[1].cells[3].text = plan.input_prompt

    cells[2].cells[0].text = "课时 / 时长"
    cells[2].cells[1].text = "—"
    cells[2].cells[2].text = "版本"
    cells[2].cells[3].text = "v1"

    doc.add_paragraph("")

    def add_section(heading: str, content: str):
        doc.add_heading(heading, level=1)
        if not content:
            doc.add_paragraph("（暂无内容）")
            return
        # 将Markdown样式的列表转为段落列表
        lines = [l.strip() for l in str(content).splitlines() if l.strip()]
        is_list_like = any(l.startswith(('-', '*', '•', '1.', '2.')) for l in lines)
        if is_list_like:
            for l in lines:
                doc.add_paragraph(l)
        else:
            doc.add_paragraph(content)

    # 各章节内容
    add_section("教学内容", get_value(key_aliases["教学内容"]))
    add_section("教学目标", get_value(key_aliases["教学目标"]))
    add_section("教学重点", get_value(key_aliases["教学重点"]))
    add_section("教学难点", get_value(key_aliases["教学难点"]))

    # 教学设计：尝试生成结构化表格
    doc.add_heading("教学设计与过程", level=1)
    design_text = get_value(key_aliases["教学设计"]).strip()
    if not design_text:
        doc.add_paragraph("（暂无内容）")
    else:
        phase_keywords = ["导入", "新授", "讲授", "练习", "巩固", "讨论", "总结", "作业"]
        rows: List[Dict[str, str]] = []
        current_phase = None
        buffer: List[str] = []
        for line in design_text.splitlines():
            l = line.strip()
            if not l:
                continue
            matched = next((p for p in phase_keywords if l.startswith(p)), None)
            if matched:
                if current_phase:
                    rows.append({"phase": current_phase, "activity": "\n".join(buffer).strip()})
                    buffer = []
                current_phase = matched
                remainder = l[len(matched):].lstrip("：: -")
                if remainder:
                    buffer.append(remainder)
            else:
                buffer.append(l)
        if current_phase:
            rows.append({"phase": current_phase, "activity": "\n".join(buffer).strip()})

        if not rows:
            # 回退为单行表格
            t = doc.add_table(rows=2, cols=2)
            t.style = 'Light Grid'
            t.cell(0, 0).text = "环节"
            t.cell(0, 1).text = "活动与要点"
            t.cell(1, 0).text = "整体"
            t.cell(1, 1).text = design_text
        else:
            t = doc.add_table(rows=1, cols=3)
            t.style = 'Light Grid'
            hdr = t.rows[0].cells
            hdr[0].text = "环节"
            hdr[1].text = "活动与要点"
            hdr[2].text = "预计时间"
            for r in rows:
                row = t.add_row().cells
                row[0].text = r.get("phase", "")
                row[1].text = r.get("activity", "")
                row[2].text = "—"

    doc.add_paragraph("")
    add_section("教学反思与总结", get_value(key_aliases["教学反思与总结"]))

    # 输出
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"教案_{plan.id}.docx"
    return StreamingResponse(
        BytesIO(buffer.read()),
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment; filename="{filename}"'}
    )


# ============== AI 连接诊断 ==============
@teacher_router.get("/ai-diagnostics")
async def ai_diagnostics(
    current_user: User = Depends(get_current_user)
):
    """测试与AI服务的连接情况，返回是否使用真实模型或回退到模拟。"""
    if current_user.role != "教师":
        raise HTTPException(status_code=403, detail="权限不足")

    try:
        resp = await ai_service.call_deepseek_api("诊断：请回复OK")
        using_mock = resp.model.startswith("mock") if hasattr(resp, 'model') else True
        return {
            "ok": True,
            "using_mock": using_mock,
            "model": getattr(resp, 'model', 'unknown'),
            "usage": getattr(resp, 'usage', {}),
            "timestamp": datetime.now()
        }
    except Exception as e:
        return {"ok": False, "error": str(e)}

# 学生疑问处理接口
@teacher_router.get("/disputes", response_model=List[StudentDisputeResponse])
async def get_student_disputes(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """获取学生疑问列表"""
    if current_user.role != "教师":
        raise HTTPException(status_code=403, detail="权限不足")
    
    # 获取教师所在班级的学生疑问
    disputes = db.query(StudentDispute).join(User).filter(
        StudentDispute.class_id == current_user.class_id
    ).order_by(StudentDispute.created_at.desc()).all()
    
    result = []
    for dispute in disputes:
        student = db.query(User).filter(User.id == dispute.student_id).first()
        result.append(StudentDisputeResponse(
            id=dispute.id,
            student_id=dispute.student_id,
            student_name=student.display_name if student else "未知学生",
            class_id=dispute.class_id,
            question_id=dispute.question_id,
            message=dispute.message,
            status=dispute.status,
            teacher_reply=dispute.teacher_reply,
            created_at=dispute.created_at,
            replied_at=dispute.replied_at
        ))
    
    return result

@teacher_router.post("/disputes/{dispute_id}/reply")
async def reply_to_dispute(
    dispute_id: int,
    reply: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """回复学生疑问"""
    if current_user.role != "教师":
        raise HTTPException(status_code=403, detail="权限不足")
    
    dispute = db.query(StudentDispute).filter(
        StudentDispute.id == dispute_id,
        StudentDispute.class_id == current_user.class_id
    ).first()
    
    if not dispute:
        raise HTTPException(status_code=404, detail="疑问不存在")
    
    dispute.teacher_reply = reply
    dispute.status = "已回复"
    dispute.replied_at = datetime.now()
    
    db.commit()
    
    return {"message": "回复成功"}

# 视频管理接口
@teacher_router.post("/videos", response_model=VideoResourceResponse)
async def create_video_resource(
    video_data: VideoResourceCreate,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """创建视频资源"""
    if current_user.role != "教师":
        raise HTTPException(status_code=403, detail="权限不足")
    
    new_video = VideoResource(
        teacher_id=current_user.id,
        title=video_data.title,
        description=video_data.description,
        path=video_data.path,
        status=video_data.status
    )
    
    db.add(new_video)
    db.commit()
    db.refresh(new_video)
    
    return VideoResourceResponse.from_orm(new_video)

@teacher_router.get("/videos", response_model=List[VideoResourceResponse])
async def get_video_resources(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """获取教师的视频资源列表"""
    if current_user.role != "教师":
        raise HTTPException(status_code=403, detail="权限不足")
    
    videos = db.query(VideoResource).filter(
        VideoResource.teacher_id == current_user.id
    ).order_by(VideoResource.created_at.desc()).all()
    
    return [VideoResourceResponse.from_orm(video) for video in videos]

@teacher_router.post("/videos/{video_id}/analyze")
async def analyze_video(
    video_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """AI分析视频内容"""
    if current_user.role != "教师":
        raise HTTPException(status_code=403, detail="权限不足")

    video = db.query(VideoResource).filter(
        VideoResource.id == video_id,
        VideoResource.teacher_id == current_user.id
    ).first()

    if not video:
        raise HTTPException(status_code=404, detail="视频不存在")

    try:
        # 调用AI服务分析视频
        analysis_result = await ai_service.analyze_video(
            video_path=video.path
        )
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"AI服务调用失败: {str(e)}。请检查API密钥配置。"
        )

    return {
        "video_id": video_id,
        "analysis": analysis_result,
        "timestamp": datetime.now()
    }
